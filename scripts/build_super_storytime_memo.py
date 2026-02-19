#!/usr/bin/env python3
import argparse, json, math, subprocess
from pathlib import Path
import pandas as pd
from docx import Document

ROOT = Path(__file__).resolve().parents[1]
OUTPUTS = ROOT / "outputs"
EXPORT  = ROOT / "export"
DATA    = ROOT / "data" / "processed"

def _isnan(x):
    return isinstance(x, float) and math.isnan(x)

def money(x):
    if x is None or _isnan(x): return "N/A"
    x = float(x)
    sgn = "-" if x < 0 else ""
    x = abs(x)
    if x >= 1e12: return f"{sgn}${x/1e12:.2f}T"
    if x >= 1e9:  return f"{sgn}${x/1e9:.2f}B"
    if x >= 1e6:  return f"{sgn}${x/1e6:.2f}M"
    return f"{sgn}${x:,.0f}"

def pct(x):
    if x is None or _isnan(x): return "N/A"
    return f"{float(x):.2f}%"

def mult(x):
    if x is None or _isnan(x): return "N/A"
    return f"{float(x):.2f}x"

def classify_rev_growth(v):
    if v is None or _isnan(v): return ("UNKNOWN", "N/A")
    v = float(v)
    if v > 10: return ("GOOD", "more than +10%")
    if v >= 0: return ("MIXED", "0% to +10%")
    return ("BAD", "below 0%")

def classify_fcf(v):
    if v is None or _isnan(v): return ("UNKNOWN", "N/A")
    v = float(v)
    if v > 0: return ("GOOD", "positive")
    return ("BAD", "negative")

def classify_fcf_margin(v):
    if v is None or _isnan(v): return ("UNKNOWN", "N/A")
    v = float(v)
    if v >= 10: return ("GOOD", "10% or higher")
    if v >= 3:  return ("MIXED", "3% to 10%")
    return ("BAD", "0% or negative")

def classify_fcf_yield(v):
    if v is None or _isnan(v): return ("UNKNOWN", "N/A")
    v = float(v)
    if v > 5:  return ("GOOD", "above 5% (often cheap)")
    if v >= 2: return ("MIXED", "2% to 5% (neutral)")
    return ("BAD", "below 2% (often expensive)")

def classify_netdebt_to_fcf(v):
    if v is None or _isnan(v): return ("UNKNOWN", "N/A")
    v = float(v)
    if v < 3: return ("GOOD", "below 3x")
    if v <= 6: return ("MIXED", "3x to 6x")
    return ("BAD", "above 6x")

def emoji(tag):
    return {"GOOD":"✅","MIXED":"🟡","BAD":"❌","UNKNOWN":"❓"}.get(tag,"❓")

def load_comps_row(ticker: str):
    p = DATA / "comps_snapshot.csv"
    if not p.exists():
        raise FileNotFoundError(f"Missing {p}")
    df = pd.read_csv(p)
    r = df[df["ticker"].astype(str).str.upper()==ticker.upper()]
    if r.empty:
        raise ValueError(f"No row for {ticker} in comps_snapshot.csv")
    row = r.iloc[0].to_dict()
    return row

def load_decision_summary(ticker: str):
    p = OUTPUTS / "decision_summary.json"
    if not p.exists(): return {}
    js = json.loads(p.read_text(encoding="utf-8"))
    if str(js.get("ticker","")).upper() != ticker.upper():
        # don’t hard fail; just return empty if mismatch
        return {}
    return js

def load_veracity(ticker: str):
    p = OUTPUTS / f"veracity_{ticker.upper()}.json"
    if not p.exists(): return {}
    return json.loads(p.read_text(encoding="utf-8"))

def load_thesis_text(thesis_path: Path, ticker: str):
    if thesis_path and thesis_path.exists():
        t = json.loads(thesis_path.read_text(encoding="utf-8"))
        headline = t.get("title") or t.get("headline") or f"{ticker}: Thesis"
        desc = t.get("description") or t.get("summary") or ""
        return headline, desc
    return f"{ticker}: Thesis", ""

def md_lines(row, ticker, thesis_headline, thesis_desc, rating, score, confidence):
    rev = row.get("revenue_ttm_yoy_pct")
    fcf = row.get("fcf_ttm")
    fcfm = row.get("fcf_margin_ttm_pct")
    # Free cash flow yield (%): prefer fcf_yield_pct, else fcf_yield*100, else compute from fcf_ttm/market_cap
    fcfy = row.get("fcf_yield_pct")
    if fcfy is None or _isnan(fcfy):
        fy = row.get("fcf_yield")
        if fy is not None and not _isnan(fy):
            try:
                fcfy = float(fy) * 100.0
            except Exception:
                pass
    if fcfy is None or _isnan(fcfy):
        try:
            mc = row.get("market_cap")
            if mc is not None and not _isnan(mc) and float(mc) != 0.0 and fcf is not None and not _isnan(fcf):
                fcfy = (float(fcf) / float(mc)) * 100.0
        except Exception:
            pass
    cash = row.get("cash")
    debt = row.get("debt")
    netd = row.get("net_debt")
    nd2f = row.get("net_debt_to_fcf_ttm")

    tag_rev, band_rev = classify_rev_growth(rev)
    tag_fcf, band_fcf = classify_fcf(fcf)
    tag_fcfm, band_fcfm = classify_fcf_margin(fcfm)
    tag_fcfy, band_fcfy = classify_fcf_yield(fcfy)
    tag_nd2f, band_nd2f = classify_netdebt_to_fcf(nd2f)

    lines = []
    lines.append(f"# SUPER Storytime Memo — {ticker}")
    lines.append(f"*Generated: {pd.Timestamp.utcnow().strftime('%Y-%m-%d %H:%M UTC')}*")
    lines.append("")
    lines.append("## 1) Your thesis (what you believe)")
    lines.append(f"**{thesis_headline}**")
    if thesis_desc.strip():
        lines.append(thesis_desc.strip())
    lines.append("")
    lines.append("## 2) What the model concluded (plain English)")
    if rating and score is not None:
        lines.append(f"- **Rating:** **{rating}** (score **{score}/100**)")
    else:
        lines.append("- **Rating:** N/A")
    lines.append(f"- **Evidence confidence:** **{confidence if confidence is not None else 'N/A'}** (higher = more trustworthy coverage)")
    lines.append("")
    lines.append("## 3) The 30-second explanation (for total beginners)")
    lines.append("Think of this like a **car dashboard**:")
    lines.append("- The **score** is overall attractiveness (health + valuation + risk).")
    lines.append("- The **numbers below** are the main gauges.")
    lines.append("- The **news/risk** is the ‘warning lights’ that can move the stock fast.")
    lines.append("")
    lines.append("## 4) Core numbers (what GM looks like today)")
    lines.append(f"- Sales growth compared to last year: **{pct(rev)}**")
    lines.append(f"- Free cash flow over the last 12 months: **{money(fcf)}**")
    lines.append(f"- Free cash flow margin (cash per $100 of sales): **{pct(fcfm)}**")
    lines.append(f"- Free cash flow yield (cash vs stock price): **{pct(fcfy)}**")
    lines.append("")
    lines.append("## 5) Good vs Bad cheat-sheet (rules + GM’s actual numbers)")
    lines.append("Each line is **(Rule) + (GM’s actual)** so the story matches the facts.")
    lines.append("")
    lines.append(f"### Sales growth compared to last year")
    lines.append(f"- {emoji(tag_rev)} **Rule band:** {band_rev}")
    lines.append(f"- **GM today:** **{pct(rev)}** → {emoji(tag_rev)} **{tag_rev}**")
    lines.append("")
    lines.append(f"### Free cash flow (cash left after bills + investment)")
    lines.append(f"- {emoji(tag_fcf)} **Rule band:** {band_fcf}")
    lines.append(f"- **GM today:** **{money(fcf)}** → {emoji(tag_fcf)} **{tag_fcf}**")
    lines.append("")
    lines.append(f"### Free cash flow margin (cash per $100 of sales)")
    lines.append(f"- {emoji(tag_fcfm)} **Rule band:** {band_fcfm}")
    lines.append(f"- **GM today:** **{pct(fcfm)}** → {emoji(tag_fcfm)} **{tag_fcfm}**")
    lines.append("")
    lines.append(f"### Free cash flow yield (cash vs what you pay for the stock)")
    lines.append(f"- {emoji(tag_fcfy)} **Rule band:** {band_fcfy}")
    lines.append(f"- **GM today:** **{pct(fcfy)}** → {emoji(tag_fcfy)} **{tag_fcfy}**")
    lines.append("")
    lines.append("### Debt snapshot (why this can bite)")
    lines.append(f"- Cash: **{money(cash)}**")
    lines.append(f"- Debt: **{money(debt)}**")
    lines.append(f"- Net debt (debt minus cash): **{money(netd)}**")
    lines.append(f"- Net debt divided by free cash flow (years to pay): **{mult(nd2f)}** → {emoji(tag_nd2f)} **{tag_nd2f}** (rule: {band_nd2f})")
    lines.append("")
    lines.append("## 6) Storytime walkthrough (explain it like I’m five)")
    lines.append("Imagine GM is a **gigantic toy factory**.")
    lines.append("")
    lines.append("### Step 1 — Are more toys being sold?")
    lines.append(f"GM’s sales growth compared to last year is **{pct(rev)}**.")
    if tag_rev == "BAD":
        lines.append("That means **fewer toys** are being sold than last year (usually a warning sign).")
    elif tag_rev == "MIXED":
        lines.append("That means sales are **flat-ish** (not exciting, but not a disaster).")
    elif tag_rev == "GOOD":
        lines.append("That means **more toys** are being sold (usually good).")
    else:
        lines.append("We don’t have enough data to judge sales growth confidently.")
    lines.append("")
    lines.append("### Step 2 — After paying for everything… is there money left?")
    lines.append(f"GM’s free cash flow over the last 12 months is **{money(fcf)}**.")
    if tag_fcf == "GOOD":
        lines.append("So the piggy bank is **filling**, not leaking. Good.")
    elif tag_fcf == "BAD":
        lines.append("So the piggy bank is **leaking**. That’s risky.")
    else:
        lines.append("We don’t have enough data to judge free cash flow confidently.")
    lines.append("")
    lines.append("### Step 3 — How efficient is the factory at turning sales into cash?")
    lines.append(f"GM’s free cash flow margin is **{pct(fcfm)}**.")
    if tag_fcfm == "GOOD":
        lines.append("That’s strong: the factory turns sales into real spare cash efficiently.")
    elif tag_fcfm == "MIXED":
        lines.append("That’s okay-ish: it makes cash, but not a ‘money printer.’")
    elif tag_fcfm == "BAD":
        lines.append("That’s weak: sales aren’t turning into much spare cash.")
    else:
        lines.append("We don’t have enough data to judge the cash margin confidently.")
    lines.append("")
    lines.append("### Step 4 — Is the stock price cheap or expensive vs cash?")
    lines.append(f"GM’s free cash flow yield is **{pct(fcfy)}**.")
    if tag_fcfy == "GOOD":
        lines.append("That often means the stock looks **cheap** relative to cash (but check if cash is sustainable).")
    elif tag_fcfy == "MIXED":
        lines.append("That’s a middle-of-the-road valuation signal.")
    elif tag_fcfy == "BAD":
        lines.append("That often means the stock looks **expensive** relative to cash.")
    else:
        lines.append("We don’t have enough data to judge valuation via cash yield confidently.")
    lines.append("")
    lines.append("### Step 5 — The big boss fight: debt")
    lines.append(f"GM has net debt of **{money(netd)}** and net debt divided by free cash flow of **{mult(nd2f)}**.")
    if tag_nd2f == "BAD":
        lines.append("That’s a **heavy backpack**. If the world gets worse, debt reduces flexibility.")
    elif tag_nd2f == "MIXED":
        lines.append("Debt is manageable but worth watching.")
    elif tag_nd2f == "GOOD":
        lines.append("Debt looks manageable relative to cash generation.")
    else:
        lines.append("We don’t have enough data to judge debt risk confidently.")
    lines.append("")
    lines.append("## 7) What to open (in the right order)")
    lines.append(f"1) Dashboard: `outputs/decision_dashboard_{ticker}.html` (speedometer + warning lights)")
    lines.append(f"2) News clickpack: `outputs/news_clickpack_{ticker}.html` (verify headline risk is real)")
    lines.append(f"3) This memo PDF: `export/{ticker}_SUPER_Storytime_Memo.pdf`")
    lines.append("")
    return "\n".join(lines)

def write_docx(md_text: str, docx_path: Path):
    doc = Document()
    for raw in md_text.splitlines():
        line = raw.rstrip()
        if line.startswith("# "):
            doc.add_heading(line[2:], level=0)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif line.startswith("- "):
            p = doc.add_paragraph(line[2:], style="List Bullet")
        elif line.startswith("1) ") or line.startswith("2) ") or line.startswith("3) "):
            doc.add_paragraph(line, style="List Number")
        elif line.strip() == "":
            doc.add_paragraph("")
        else:
            doc.add_paragraph(line)
    doc.save(docx_path)

def export_pdf(docx_path: Path, pdf_path: Path):
    soffice = Path("/opt/homebrew/bin/soffice")
    if not soffice.exists():
        return False
    cmd = [str(soffice), "--headless", "--convert-to", "pdf", "--outdir", str(pdf_path.parent), str(docx_path)]
    subprocess.run(cmd, check=False)
    return pdf_path.exists()

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--ticker", required=True)
    ap.add_argument("--thesis", required=False, default="")
    args = ap.parse_args()

    ticker = args.ticker.upper()
    thesis_path = Path(args.thesis) if args.thesis else None

    row = load_comps_row(ticker)
    ds  = load_decision_summary(ticker)
    ver = load_veracity(ticker)

    rating = ds.get("rating")
    score  = ds.get("score")
    confidence = ver.get("confidence") or ver.get("veracity") or ds.get("confidence")

    th_head, th_desc = load_thesis_text(thesis_path, ticker)

    md_text = md_lines(row, ticker, th_head, th_desc, rating, score, confidence)

    OUTPUTS.mkdir(parents=True, exist_ok=True)
    EXPORT.mkdir(parents=True, exist_ok=True)

    md_path = OUTPUTS / f"{ticker}_SUPER_Storytime_Memo.md"
    docx_path = EXPORT / f"{ticker}_SUPER_Storytime_Memo.docx"
    pdf_path  = EXPORT / f"{ticker}_SUPER_Storytime_Memo.pdf"

    md_path.write_text(md_text, encoding="utf-8")
    write_docx(md_text, docx_path)
    export_pdf(docx_path, pdf_path)

    print("DONE ✅ SUPER Storytime memo created:")
    print(f"- {md_path}")
    print(f"- {docx_path}")
    if pdf_path.exists():
        print(f"- {pdf_path}")
    else:
        print("⚠️ PDF export failed (soffice missing or conversion failed).")

if __name__ == "__main__":
    main()
