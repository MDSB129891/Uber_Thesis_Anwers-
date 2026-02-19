#!/usr/bin/env python3
import argparse, json, math, subprocess
from pathlib import Path
from datetime import datetime, timezone

import pandas as pd
from docx import Document

REPO = Path(__file__).resolve().parents[1]
DATA = REPO / "data" / "processed"
OUT = REPO / "outputs"
EXP = REPO / "export"

def money(x):
    if x is None or (isinstance(x, float) and math.isnan(x)):
        return "N/A"
    ax = abs(float(x))
    if ax >= 1e12: return f"${x/1e12:,.2f}T"
    if ax >= 1e9:  return f"${x/1e9:,.2f}B"
    if ax >= 1e6:  return f"${x/1e6:,.2f}M"
    return f"${x:,.0f}"

def pct(x):
    if x is None or (isinstance(x, float) and math.isnan(x)):
        return "N/A"
    return f"{100*float(x):.2f}%"

def safe_float(v):
    try:
        if v is None: return None
        if isinstance(v, str) and not v.strip(): return None
        f = float(v)
        if math.isnan(f): return None
        return f
    except Exception:
        return None

def dcf_fcff(rev0, rev_cagr, fcf_margin, years, wacc, terminal_g):
    # Project revenue and FCF directly (simple + stable)
    rev = []
    fcf = []
    for t in range(1, years+1):
        rt = rev0 * ((1.0 + rev_cagr) ** t)
        rev.append(rt)
        fcf.append(rt * fcf_margin)

    pv = 0.0
    for t, cf in enumerate(fcf, start=1):
        pv += cf / ((1.0 + wacc) ** t)

    # Terminal value at end of year N
    cfN = fcf[-1]
    if wacc <= terminal_g:
        tv = None
    else:
        tv = (cfN * (1.0 + terminal_g)) / (wacc - terminal_g)

    if tv is None:
        return {"pv_cf": pv, "tv": None, "pv_tv": None, "ev": None}
    pv_tv = tv / ((1.0 + wacc) ** years)
    ev = pv + pv_tv
    return {"pv_cf": pv, "tv": tv, "pv_tv": pv_tv, "ev": ev}

def render_sensitivity(rev0, fcf_margin, years, wacc_grid, g_grid):
    rows = []
    for w in wacc_grid:
        row = {"WACC": pct(w)}
        for g in g_grid:
            r = dcf_fcff(rev0, 0.10, fcf_margin, years, w, g)  # use base rev_cagr for grid
            row[f"g={pct(g)}"] = money(r["ev"]) if r["ev"] is not None else "N/A"
        rows.append(row)
    return pd.DataFrame(rows)

def doc_add_h1(doc, s): doc.add_heading(s, level=1)
def doc_add_h2(doc, s): doc.add_heading(s, level=2)
def doc_add_p(doc, s): doc.add_paragraph(s)

def main(ticker: str, assumptions_path: Path):
    ticker = ticker.upper().strip()
    snap_path = DATA / "comps_snapshot.csv"
    if not snap_path.exists():
        raise SystemExit(f"Missing {snap_path}. Run your engine update first.")

    df = pd.read_csv(snap_path)
    df["ticker"] = df["ticker"].astype(str).str.upper()
    r = df[df["ticker"] == ticker]
    if r.empty:
        raise SystemExit(f"Ticker {ticker} not found in {snap_path}")

    row = r.iloc[0].to_dict()

    # Inputs from your pipeline
    rev0 = safe_float(row.get("revenue_ttm"))
    fcf_margin = safe_float(row.get("fcf_margin_ttm_pct"))
    if fcf_margin is not None:
        fcf_margin = fcf_margin / 100.0

    market_cap = safe_float(row.get("market_cap"))
    cash = safe_float(row.get("cash"))
    debt = safe_float(row.get("debt"))
    net_debt = safe_float(row.get("net_debt"))
    if net_debt is None and (cash is not None and debt is not None):
        net_debt = debt - cash

    if rev0 is None or fcf_margin is None:
        raise SystemExit(
            "Not enough data for DCF (need revenue_ttm + fcf_margin_ttm_pct). "
            "Check data/processed/comps_snapshot.csv for this ticker."
        )

    a = json.loads(assumptions_path.read_text(encoding="utf-8"))
    years = int(a.get("projection_years", 5))
    scenarios = a["scenarios"]
    wacc_grid = a.get("wacc_grid", [0.085, 0.095, 0.105])
    g_grid = a.get("terminal_g_grid", [0.02, 0.025, 0.03])

    now = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")

    # Run scenarios
    scen_out = {}
    for name, s in scenarios.items():
        out = dcf_fcff(
            rev0=rev0,
            rev_cagr=float(s["rev_cagr"]),
            fcf_margin=float(s["fcf_margin"]),
            years=years,
            wacc=float(s["wacc"]),
            terminal_g=float(s["terminal_g"])
        )
        ev = out["ev"]
        eq = None if ev is None else (ev - (net_debt or 0.0))
        scen_out[name] = {**out, "equity_value": eq}

    # Sensitivity grid (EV only)
    sens = render_sensitivity(rev0, fcf_margin, years, wacc_grid, g_grid)

    # Write MD
    md_path = OUT / f"{ticker}_DCF_Appendix.md"
    lines = []
    lines.append(f"# DCF Appendix — {ticker}")
    lines.append(f"*Generated: {now}*")
    lines.append("")
    lines.append("## What this is (for the skeptics)")
    lines.append("This is a **purely numbers-based** discounted cash flow (DCF) appendix.")
    lines.append("It uses your pipeline’s **Revenue (last 12 months)** and **Free cash flow margin** to project cash flows, then discounts them back using WACC.")
    lines.append("")
    lines.append("## Inputs used (traceable)")
    lines.append(f"- Revenue (last 12 months): **{money(rev0)}**  _(source: comps_snapshot → revenue_ttm)_")
    lines.append(f"- Free cash flow margin (last 12 months): **{pct(fcf_margin)}**  _(source: comps_snapshot → fcf_margin_ttm_pct)_")
    lines.append(f"- Market cap: **{money(market_cap)}**  _(source: comps_snapshot → market_cap)_")
    lines.append(f"- Net debt: **{money(net_debt)}**  _(source: comps_snapshot → net_debt or debt−cash)_")
    lines.append("")
    lines.append("## Scenario results (enterprise value and equity value)")
    lines.append("| Scenario | Revenue growth (per year) | Free cash flow margin | WACC | Terminal growth | Enterprise value | Equity value |")
    lines.append("|---|---:|---:|---:|---:|---:|---:|")
    for name in ["bear","base","bull"]:
        s = scenarios[name]
        o = scen_out[name]
        lines.append(
            f"| {name.upper()} | {pct(s['rev_cagr'])} | {pct(s['fcf_margin'])} | {pct(s['wacc'])} | {pct(s['terminal_g'])} | {money(o['ev'])} | {money(o['equity_value'])} |"
        )
    lines.append("")
    lines.append("## Sensitivity (enterprise value) — WACC × Terminal growth")
    lines.append("")
    lines.append(sens.to_markdown(index=False))
    lines.append("")
    lines.append("## Important honesty (so nobody over-trusts it)")
    lines.append("- If the business gets hit by a **cost shock** (ex: drivers become employees), the key DCF levers are **free cash flow margin** and **WACC** (risk).")
    lines.append("- DCF is not a truth machine. It’s a calculator: **garbage in → garbage out**, so we show the full sensitivity grid.")
    lines.append("")
    md_path.write_text("\n".join(lines), encoding="utf-8")

    # Write DOCX (simple but clean)
    docx_path = EXP / f"{ticker}_DCF_Appendix.docx"
    doc = Document()
    doc_add_h1(doc, f"DCF Appendix — {ticker}")
    doc_add_p(doc, f"Generated: {now}")

    doc_add_h2(doc, "What this is (for the skeptics)")
    doc_add_p(doc, "This is a purely numbers-based discounted cash flow (DCF) appendix.")
    doc_add_p(doc, "It uses Revenue (last 12 months) and Free cash flow margin to project cash flows, then discounts them back using WACC.")

    doc_add_h2(doc, "Inputs used (traceable)")
    doc_add_p(doc, f"Revenue (last 12 months): {money(rev0)} (source: comps_snapshot → revenue_ttm)")
    doc_add_p(doc, f"Free cash flow margin (last 12 months): {pct(fcf_margin)} (source: comps_snapshot → fcf_margin_ttm_pct)")
    doc_add_p(doc, f"Market cap: {money(market_cap)} (source: comps_snapshot → market_cap)")
    doc_add_p(doc, f"Net debt: {money(net_debt)} (source: comps_snapshot → net_debt or debt−cash)")

    doc_add_h2(doc, "Scenario results")
    table = doc.add_table(rows=1, cols=7)
    hdr = table.rows[0].cells
    hdr[0].text = "Scenario"
    hdr[1].text = "Revenue growth / year"
    hdr[2].text = "Free cash flow margin"
    hdr[3].text = "WACC"
    hdr[4].text = "Terminal growth"
    hdr[5].text = "Enterprise value"
    hdr[6].text = "Equity value"

    for name in ["bear","base","bull"]:
        s = scenarios[name]
        o = scen_out[name]
        rowc = table.add_row().cells
        rowc[0].text = name.upper()
        rowc[1].text = pct(s["rev_cagr"])
        rowc[2].text = pct(s["fcf_margin"])
        rowc[3].text = pct(s["wacc"])
        rowc[4].text = pct(s["terminal_g"])
        rowc[5].text = money(o["ev"])
        rowc[6].text = money(o["equity_value"])

    doc_add_h2(doc, "Sensitivity (enterprise value) — WACC × Terminal growth")
    sens_table = doc.add_table(rows=1, cols=1 + len(g_grid))
    h = sens_table.rows[0].cells
    h[0].text = "WACC"
    for j, g in enumerate(g_grid, start=1):
        h[j].text = f"g={pct(g)}"

    for _, rr in sens.iterrows():
        cells = sens_table.add_row().cells
        cells[0].text = str(rr["WACC"])
        for j, g in enumerate(g_grid, start=1):
            cells[j].text = str(rr[f"g={pct(g)}"])

    doc_add_h2(doc, "Important honesty (so nobody over-trusts it)")
    doc_add_p(doc, "If the business gets hit by a cost shock (example: drivers become employees), the key DCF levers are free cash flow margin and WACC (risk).")
    doc_add_p(doc, "DCF is not a truth machine. It’s a calculator: garbage in → garbage out, so we show the full sensitivity grid.")

    doc.save(docx_path)

    # Export PDF via soffice if available
    pdf_path = EXP / f"{ticker}_DCF_Appendix.pdf"
    soffice = None
    for c in ["soffice", "/opt/homebrew/bin/soffice"]:
        if Path(c).exists() or (c == "soffice"):
            soffice = c
            break

    if soffice:
        try:
            subprocess.run([soffice, "--headless", "--convert-to", "pdf", "--outdir", str(EXP), str(docx_path)], check=False)
        except Exception:
            pass

    print("DONE ✅ DCF appendix created:")
    print(f"- {md_path}")
    print(f"- {docx_path}")
    if pdf_path.exists():
        print(f"- {pdf_path}")
    else:
        print("⚠️ PDF not created automatically. Try:")
        print(f"  /opt/homebrew/bin/soffice --headless --convert-to pdf --outdir export {docx_path}")

if __name__ == "__main__":
    ap = argparse.ArgumentParser()
    ap.add_argument("--ticker", required=True)
    ap.add_argument("--assumptions", default="dcf_assumptions/default.json")
    args = ap.parse_args()

    main(args.ticker, Path(args.assumptions))
