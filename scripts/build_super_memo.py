STORY_BLOCK = r'''
## Good vs Bad cheat-sheet (how to judge the numbers)


### 1) Revenue growth (year-over-year)
**What it is:** “Are sales bigger than last year?”
- ✅ **Usually good:** **> +10%** (strong growth)
- 🟡 **Okay / depends:** **0% to +10%**
- ❌ **Usually bad:** **< 0%** (shrinking sales)  
**Why it matters:** shrinking sales often means demand is weakening.

### 2) Free cash flow
**What it is:** “After paying all bills AND investing in the business, is there cash left?”
- ✅ **Good:** **positive** and stable/increasing
- 🟡 **Mixed:** small positive but volatile (up/down a lot)
- ❌ **Bad:** negative consistently (burning cash)  
**Why it matters:** cash pays debt, buybacks, dividends, growth, and survival.

### 3) Free cash flow margin
**What it is:** “Out of $100 of sales, how many dollars become free cash?”
- ✅ **Good:** **10%+**
- 🟡 **Okay:** **3%–10%**
- ❌ **Bad:** **0% or negative**  
**Why it matters:** a company can have big revenue but still be a cash-bleeder.

### 4) Free cash flow yield (cash return vs stock price)
**What it is:** “How much cash does the business generate compared to what you pay for the stock?”
- ✅ **Often good/cheap:** **> 5%**
- 🟡 **Neutral:** **2%–5%**
- ❌ **Often expensive:** **< 2%**  
**Why it matters:** yields can be high because the market expects trouble — so always cross-check.

### 5) Net debt (debt minus cash)
**What it is:** “If the company used all its cash to pay debt, what debt is left?”
- ✅ **Better:** low net debt (or net cash)
- 🟡 **Okay:** moderate net debt if cash flow is strong
- ❌ **Risky:** huge net debt with weakening cash flow  
**Why it matters:** debt is a fixed obligation. Weak cash flow + big debt can spiral.

### 6) Net debt / free cash flow (years to pay debt)
**What it is:** “How many years of today’s free cash flow to pay off net debt?”
- ✅ **Good:** **< 3x**
- 🟡 **Watch:** **3x–6x**
- ❌ **High risk:** **> 6x**

### 7) News shock (last 30 days)
**What it is:** “Did headlines recently turn sharply negative?”
- ✅ **Good:** mild/normal (no crisis)
- 🟡 **Watch:** elevated negativity (verify details)
- ❌ **Bad:** severe negative burst (often hits stock fast)

### 8) Risk buckets (labor / regulatory / insurance)
**What it is:** “Is the company being hit with frequent negative headlines about these risks?”
- ✅ **Good:** low frequency
- 🟡 **Watch:** moderate frequency
- ❌ **Bad:** high frequency / increasing

---

## Storytime walkthrough (explain it like I’m five)

### Step A — Sales (Revenue)

### Step B — Real cash (Free cash flow)

### Step C — Efficiency (Free cash flow margin)

### Step D — Price vs cash (Free cash flow yield)

---

## Bucket storytime (what each bucket means + how to improve it)


### cash_level (Cash Level)
**Meaning:** “Does the business create real cash and have breathing room?”
- Better when: free cash flow is positive and stable; cash is healthy.
- Worse when: cash flow is negative/weak or falling.

### valuation (Price vs Reality)
**Meaning:** “Are you paying a fair price for the cash the business produces?”
- Better when: free cash flow yield is higher (often cheaper).
- Worse when: free cash flow yield is very low (often expensive) OR cash flow is questionable.

### growth (Is the business expanding?)
**Meaning:** “Are sales and/or cash growing?”
- Better when: revenue and free cash flow are growing versus last year.
- Worse when: revenue shrinks or free cash flow shrinks (especially both).

### quality (Business health)
**Meaning:** “Is this a good machine or a fragile machine?”
- Better when: margins are stable and results are consistent.
- Worse when: margins are volatile and surprises happen often.

### balance_risk (Debt / blow-up risk)
**Meaning:** “Could debt + obligations cause a faceplant if conditions worsen?”
- Better when: net debt is low.
- Risky when: net debt is huge and cash flow is weakening.
'''

#!/usr/bin/env python3
import argparse, json

def _de_jargon(s: str) -> str:
    if not isinstance(s, str) or not s:
        return s

    # Normalize weird Word bullets into normal markdown bullets
    s = s.replace(" ", "- ")

    # Expand abbreviations (aggressive on purpose)
    s = re.sub(r"\bcompared to last year\b", "compared to the same time last year", s)
    s = re.sub(r"\(compared to last year\)", "(compared to the same time last year)", s)

    s = re.sub(r"\bover the last 12 months\b", "over the last twelve months", s)
    s = re.sub(r"\(over the last 12 months\)", "(over the last twelve months)", s)

    s = re.sub(r"\bfree cash flow\b", "free cash flow", s)
    s = s.replace("free cash flow margin", "free cash flow margin")
    s = s.replace("free cash flow yield", "free cash flow yield")

    # Common red-flag phrases
    s = s.replace("over the last 12 months revenue declining compared to last year", "Revenue is shrinking compared to last year (last twelve months)")
    s = s.replace("over the last 12 months free cash flow declining compared to last year", "Free cash flow is shrinking compared to last year (last twelve months)")
    s = s.replace("Net debt high vs over the last 12 months free cash flow", "Debt looks heavy compared to cash the business generates (last twelve months)")

    # Bucket names (optional nicer words)
    s = s.replace("**cash_level**", "**Cash strength**")
    s = s.replace("**valuation**", "**Price vs value**")
    s = s.replace("**growth**", "**Growth**")
    s = s.replace("**quality**", "**Business quality**")
    s = s.replace("**balance_risk**", "**Debt / balance-sheet risk**")

    return s


from pathlib import Path
from datetime import datetime
import pandas as pd
from docx import Document

import re

def humanize_terms(s: str) -> str:
    """
    """
    if not isinstance(s, str) or not s:
        return s

    # Whole-word replacements (case-sensitive where it matters)
    repl = [
        (r"\bcompared to last year\b", "year over year (compared to the same time last year)"),
        (r"\bover the last 12 months\b", "the last twelve months"),
        (r"\bfree cash flow\b", "free cash flow (cash left after paying bills and investing in the business)"),
        (r"\bfree cash flow margin\b", "free cash flow margin (cash left from each $100 of sales)"),
        (r"\bfree cash flow yield\b", "free cash flow yield (cash return compared to the stock price)"),
        (r"\bEV\b", "enterprise value (company value including debt)"),
        (r"\bEBITDA\b", "earnings before interest, taxes, depreciation, and amortization"),
        (r"\bEPS\b", "earnings per share"),
        (r"\bP/E\b", "price-to-earnings ratio"),
        (r"\bNet debt\b", "net debt (total debt minus cash)"),
        (r"\bcapex\b", "capital spending (money used to buy/build long-term assets)"),
        (r"\bCAGR\b", "compound annual growth rate"),
        (r"\bbps\b", "basis points (one hundredth of a percent)"),

    out = s
    for pat, rep in repl:
        out = re.sub(pat, rep, out)

    # Optional polish: if something writes "over the last 12 months revenue" etc
    out = re.sub(r"\bthe last twelve months revenue\b", "revenue over the last twelve months", out)
    return out

ROOT = Path(__file__).resolve().parents[1]
OUTPUTS = ROOT / "outputs"
EXPORT = ROOT / "export"
DATA = ROOT / "data" / "processed"

def utc_now():
    return datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC")

def is_na(x):
    return x is None or (isinstance(x, float) and pd.isna(x))

def money(x):
    if is_na(x): return "N/A"
    x = float(x)
    if abs(x) >= 1e9:  return f"${x/1e9:,.2f}B"
    if abs(x) >= 1e6:  return f"${x/1e6:,.2f}M"
    return f"${x:,.0f}"

def pct(x):
    if is_na(x): return "N/A"
    return f"{float(x):.2f}%"

def num(x):
    if is_na(x): return "N/A"
    try:
        return f"{float(x):,.2f}"
    except:
        return str(x)

def load_json(p: Path):
    if not p.exists(): return None
    return json.loads(p.read_text(encoding="utf-8"))

def safe_read_csv(p: Path):
    if not p.exists(): return None
    return pd.read_csv(p)

def build_metric_lookup(ticker: str):
    T = ticker.upper()
    out = {}

    # 1) decision_summary.json (already “what engine used”)
    ds = load_json(OUTPUTS / "decision_summary.json") or {}
    ki = (ds.get("key_inputs_used", {}) or {}).copy()

    # 2) comps_snapshot.csv (valuation + over the last 12 months metrics)
    comps = safe_read_csv(DATA / "comps_snapshot.csv")
    if comps is not None and "ticker" in comps.columns:
        r = comps[comps["ticker"].astype(str).str.upper() == T]
        if not r.empty:
            row = r.iloc[0].to_dict()
            # preferred set
            for k in ["revenue_ttm_yoy_pct","fcf_ttm","fcf_margin_ttm_pct","cash","debt","net_debt","market_cap","fcf_yield_pct","fcf_yield","net_debt_to_fcf_ttm"]:
                if k in row:
                    out[k] = row.get(k)

            # fallback: some builds store yield as decimal in fcf_yield
            if is_na(out.get("fcf_yield_pct")) and not is_na(out.get("fcf_yield")):
                try:
                    out["fcf_yield_pct"] = float(out["fcf_yield"]) * 100.0
                except:
                    pass

            # map “latest_*” to over the last 12 months when needed (so claims don’t go N/A)
            if is_na(out.get("latest_revenue_yoy_pct")) and not is_na(row.get("revenue_ttm_yoy_pct")):
                out["latest_revenue_yoy_pct"] = row.get("revenue_ttm_yoy_pct")
            if is_na(out.get("latest_free_cash_flow")) and not is_na(row.get("fcf_ttm")):
                out["latest_free_cash_flow"] = row.get("fcf_ttm")
            if is_na(out.get("latest_fcf_margin_pct")) and not is_na(row.get("fcf_margin_ttm_pct")):
                out["latest_fcf_margin_pct"] = row.get("fcf_margin_ttm_pct")

    # 3) news_sentiment_proxy.csv (shock/neg counts)
    proxy = safe_read_csv(DATA / "news_sentiment_proxy.csv")
    if proxy is not None and "ticker" in proxy.columns:
        r = proxy[proxy["ticker"].astype(str).str.upper() == T]
        if not r.empty:
            pr = r.iloc[0].to_dict()
            out["news_shock_7d"] = pr.get("shock_7d")
            out["news_shock_30d"] = pr.get("shock_30d")
            out["news_neg_7d"] = pr.get("neg_7d")
            out["news_neg_30d"] = pr.get("neg_30d")
            out["news_articles_7d"] = pr.get("articles_7d")
            out["news_articles_30d"] = pr.get("articles_30d")

    # 4) risk dashboard (if exists)
    risk = safe_read_csv(DATA / "news_risk_dashboard.csv")
    if risk is not None and "ticker" in risk.columns:
        r = risk[risk["ticker"].astype(str).str.upper() == T]
        if not r.empty:
            rr = r.iloc[0].to_dict()
            for k,v in rr.items():
                if isinstance(k,str) and k.lower().startswith("risk_"):
                    out[k] = v

    return out

def eval_claim(value, op, threshold):
    if is_na(value): return "UNKNOWN"
    try:
        v = float(value)
        t = float(threshold)
    except:
        return "UNKNOWN"
    if op == ">=": return "PASS" if v >= t else "FAIL"
    if op == ">":  return "PASS" if v >  t else "FAIL"
    if op == "<=": return "PASS" if v <= t else "FAIL"
    if op == "<":  return "PASS" if v <  t else "FAIL"
    if op in ("==","="): return "PASS" if v == t else "FAIL"
    return "UNKNOWN"

def bucket_explain():
    return {
        "cash_level": "Cash Level = does the business generate real cash and have liquidity?",
        "valuation": "Valuation = are you paying a reasonable price vs the cash the business produces?",
        "growth": "Growth = are sales/cash expanding or shrinking?",
        "quality": "Quality = is the business healthy (margins, stability, consistency)?",
        "balance_risk": "Balance Risk = debt + leverage + anything that can blow up fast."

def build_story_layer(m):
    out = []

    rg = float(m.get("revenue_yoy", 0) or 0)
    fcf = float(m.get("fcf", 0) or 0)
    margin = float(m.get("fcf_margin", 0) or 0)
    debt = float(m.get("net_debt", 0) or 0)
    dy = float(m.get("fcf_yield", 0) or 0)

    if rg < 0:
        pass
    else:

        pass
    if margin < 10:
        pass
    if fcf < 0:

        pass
    out.append("Debt is future money already spent. More debt = less flexibility.\n")



    return "\n".join(out)




def _super_postprocess_text(s: str) -> str:
    """
    """
    if not isinstance(s, str) or not s:
        return s

    # Make common labels beginner-friendly
    s = s.replace("(compared to last year)", "(compared to the same time last year)")
    s = s.replace("(over the last 12 months)", "(over the last twelve months)")
    s = s.replace("compared to last year", "year over year (compared to last year)")
    s = s.replace("over the last 12 months", "the last twelve months")

    # Expand common abbreviations
    s = re.sub(r"\bfree cash flow\b", "free cash flow (cash left after paying bills and investing in the business)", s)
    s = re.sub(r"\bfree cash flow margin\b", "free cash flow margin (cash left from each $100 of sales)", s)
    s = re.sub(r"\bfree cash flow yield\b", "free cash flow yield (cash return compared to the stock price)", s)
    s = re.sub(r"\bNet debt\b", "net debt (total debt minus cash)", s)

    # Expand a few finance acronyms if they appear
    s = re.sub(r"\bEPS\b", "earnings per share", s)
    s = re.sub(r"\bP/E\b", "price-to-earnings ratio", s)
    s = re.sub(r"\bEV\b", "enterprise value (company value including debt)", s)
    s = re.sub(r"\bEBITDA\b", "earnings before interest, taxes, depreciation, and amortization", s)

    # Make bucket labels nicer
    s = s.replace("**cash_level**", "**Cash strength**")
    s = s.replace("**valuation**", "**Price vs value**")
    s = s.replace("**growth**", "**Growth**")
    s = s.replace("**quality**", "**Business quality**")
    s = s.replace("**balance_risk**", "**Debt / balance-sheet risk**")

    return s



import subprocess
from pathlib import Path

def _export_super_pdf(docx_path: Path) -> Path:
    """
    """
    docx_path = Path(docx_path)
    outdir = docx_path.parent
    pdf_path = outdir / (docx_path.stem + ".pdf")

    # Try common soffice locations
    candidates = [
        "soffice",
        "/opt/homebrew/bin/soffice",
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",

    soffice = None
    for c in candidates:
        try:
            r = subprocess.run([c, "--version"], capture_output=True, text=True)
            if r.returncode == 0:
                soffice = c
                break
        except Exception:
            pass

    if not soffice:

        [soffice, "--headless", "--convert-to", "pdf", "--outdir", str(outdir), str(docx_path)],
        check=True
    return pdf_path




STORYTIME_WALKTHROUGH_BLOCK = """
## Good vs Bad cheat-sheet (how to judge the numbers)


### 1) Revenue growth (year over year)
**What it is:** “Are sales bigger than last year?”
- ✅ **Usually good:** **> +10%** (strong growth)
- 🟡 **Okay / depends:** **0% to +10%**
- ❌ **Usually bad:** **< 0%** (shrinking sales)
**Why it matters:** shrinking sales often means demand is weakening.

### 2) Free cash flow
**What it is:** “After paying all bills AND investing in the business, is there cash left?”
- ✅ **Good:** positive and stable/increasing
- 🟡 **Mixed:** small positive but volatile (up/down a lot)
- ❌ **Bad:** negative consistently (burning cash)
**Why it matters:** cash pays debt, buybacks, dividends, growth, and survival.

### 3) Free cash flow margin
**What it is:** “Out of $100 of sales, how many dollars become free cash?”
- ✅ **Good:** **10%+** (strong cash conversion in many industries)
- 🟡 **Okay:** **3%–10%** (depends on industry)
- ❌ **Bad:** **0% or negative**
**Why it matters:** a company can have big revenue but still bleed cash.

### 4) Free cash flow yield (cash return vs stock price)
**What it is:** “How much cash does the business generate compared to what you pay for the stock?”
- ✅ **Often good/cheap:** **> 5%**
- 🟡 **Neutral:** **2%–5%**
- ❌ **Often expensive:** **< 2%**
**Why it matters:** it’s a “cash version” of valuation — but yields can be high because the market expects trouble.

### 5) Net debt (debt minus cash)
**What it is:** “If the company used all its cash to pay debt, what debt is left?”
- ✅ Better: low net debt (or net cash)
- 🟡 Okay: moderate net debt if cash flow is strong
- ❌ Risky: huge net debt with weakening cash flow
**Why it matters:** debt is a fixed obligation.

### 6) Net debt / free cash flow (years to pay debt)
**What it is:** “How many years of today’s free cash flow to pay off net debt?”
- ✅ Good: **< 3x**
- 🟡 Watch: **3x–6x**
- ❌ High risk: **> 6x**
**Why it matters:** higher = less flexibility in a downturn.

### 7) News shock (last 30 days)
**What it is:** “Did headlines recently turn sharply negative?”
- ✅ Good: mild/normal (no crisis)
- 🟡 Watch: elevated negativity (verify details)
- ❌ Bad: severe negative burst (often hits stock fast)

### 8) Risk buckets (labor / regulatory / insurance)
**What it is:** “Is the company getting hit with frequent negative headlines here?”
- ✅ Good: low frequency
- 🟡 Watch: moderate frequency
- ❌ Bad: high frequency / increasing

---

## Storytime walkthrough (explain it like I’m five)

### Step A — Sales (revenue)

### Step B — Real cash (free cash flow)
Positive = piggy bank fills. Negative = piggy bank leaks.

### Step C — Efficiency (free cash flow margin)

### Step D — Price vs cash (free cash flow yield)

---

## Bucket storytime (what each bucket means)

Each bucket is a stat. Higher stat = better chances of winning.

### Cash Level

### Valuation

### Growth

### Quality

### Balance Risk

---

## Rule of thumb (for beginners)
"""


def main(ticker: str, thesis_path: str):
    md = []  # Stormbreaker: memo lines accumulator

    T = ticker.upper()
    thesis_p = Path(thesis_path)
    thesis = json.loads(thesis_p.read_text(encoding="utf-8")) if thesis_p.exists() else {}

    ds = load_json(OUTPUTS / "decision_summary.json") or {}
    ver = load_json(OUTPUTS / f"veracity_{T}.json") or {}
    alerts = load_json(OUTPUTS / f"alerts_{T}.json") or {}
    claim_pack = load_json(OUTPUTS / f"claim_evidence_{T}.json") or {}

    metrics = build_metric_lookup(T)

    md.append(build_story_layer(metrics))

    score = ds.get("score")
    rating = ds.get("rating")
    buckets = ds.get("bucket_scores", {}) or {}
    red_flags = ds.get("red_flags", []) or []

    claims = thesis.get("claims", []) or []

    # Build claims results
    claim_rows = []
    for c in claims:
        m = c.get("metric")
        op = c.get("operator")
        th = c.get("threshold")
        val = metrics.get(m)
        status = eval_claim(val, op, th)
            "status": status,
            "statement": c.get("statement",""),
            "metric": m,
            "op": op,
            "threshold": th,
            "value": val,
            "weight": c.get("weight",1),
        })

    # Write Markdown SUPER memo
    md = []
    md.append(f"# SUPER Investment Memo — {T}")
    md.append(f"*Generated: {utc_now()}*")
    md.append("")
    md.append("## 1) Your thesis (what you believe)")
    md.append(f"**{thesis.get('name','(no name)')}**")
    md.append(thesis.get("description","(no description)"))
    md.append("")
    md.append("## 2) What the model concluded (plain English)")
    md.append(f"- **Rating:** **{rating}** (score **{score}/100**)")
    conf = ver.get("confidence") or ver.get("veracity") or ver.get("confidence_score")
    md.append(f"- **Evidence confidence / veracity:** **{conf if conf is not None else 'N/A'}** (higher = more trustworthy coverage)")
    md.append("")
    md.append("## 3) The 30-second explanation (for total beginners)")
    md.append("Think of this like a **car dashboard**:")
    md.append("- The **score** is the overall attractiveness estimate.")
    md.append("- The **buckets** explain *why* the score happened.")
    md.append("- The **news/risk** items try to spot headline landmines.")
    md.append("- The **thesis test** checks whether the facts match the story you’re betting on.")
    md.append("")

    md.append("## 4) Core numbers (sanity-check)")
    md.append(f"- Revenue growth (compared to last year): **{pct(metrics.get('latest_revenue_yoy_pct'))}**  _(source: comps_snapshot → revenue_ttm_yoy_pct)_")
    md.append(f"- Free cash flow (over the last 12 months): **{money(metrics.get('latest_free_cash_flow'))}**  _(source: comps_snapshot → fcf_ttm)_")
    md.append(f"- free cash flow margin: **{pct(metrics.get('latest_fcf_margin_pct'))}**  _(source: comps_snapshot → fcf_margin_ttm_pct)_")
    md.append(f"- free cash flow yield: **{pct(metrics.get('fcf_yield_pct'))}**  _(source: comps_snapshot → fcf_yield_pct / fcf_yield)_")
    md.append("")

    md.append("## 5) Balance sheet snapshot (why debt matters)")
    md.append(f"- Market cap: **{money(metrics.get('market_cap'))}**")
    md.append(f"- Cash: **{money(metrics.get('cash'))}**")
    md.append(f"- Debt: **{money(metrics.get('debt'))}**")
    md.append(f"- Net debt: **{money(metrics.get('net_debt'))}**  _(debt minus cash)_")
    md.append(f"- Net debt / free cash flow: **{num(metrics.get('net_debt_to_fcf_ttm'))}x**  _(how many years of cash it takes to pay debt)_")
    md.append("")

    md.append("## 6) Bucket scores (what drove the rating)")
    expl = bucket_explain()
    for k in ["cash_level","valuation","growth","quality","balance_risk"]:
        md.append(f"- **{k}** = **{buckets.get(k,'N/A')}** → {expl.get(k,'')}")
    md.append("")

    md.append("## 7) Red flags (things that can hurt stock fast)")


## What these red flags actually mean (plain English)

### 1) Revenue shrinking over the last 12 months

### 2) Free cash flow shrinking over the last 12 months


Declining free cash flow means:

- less money to pay debt
- less money to invest in electric vehicles
- less money for buybacks
- less margin for mistakes


---

### 3) High debt compared to cash generation





---

### 4) Lots of negative labor / insurance / regulatory headlines

In the last 30 days there were many negative stories about:

- workers and unions
- insurance costs
- government regulation



---

### Simple rule for beginners





## What these red flags actually mean (plain English)

### 1) Revenue shrinking over the last 12 months



    if red_flags:
        for rf in red_flags:
            md.append(f"- {rf}")
    else:
        md.append("- None flagged by the engine.")
    md.append("")

    md.append("## 8) Thesis test (PASS/FAIL vs your claims)")
    if claim_rows:
        for r in claim_rows:
            v = r["value"]
            actual = pct(v) if "pct" in (r["metric"] or "") or "margin" in (r["metric"] or "") or "yoy" in (r["metric"] or "") or "yield" in (r["metric"] or "") else money(v) if "cash" in (r["metric"] or "") or "fcf" in (r["metric"] or "") else num(v)
            md.append(f"- **{r['status']}** — {r['statement']}  \n  Metric `{r['metric']}` {r['op']} {r['threshold']} | Actual: **{actual}**")
    else:
        md.append("- No claims found in thesis file.")
    md.append("")

    md.append("## 9) What to open (dopamine mode)")
    md.append(f"- Dashboard: `outputs/decision_dashboard_{T}.html`")
    md.append(f"- News clickpack: `outputs/news_clickpack_{T}.html`")
    md.append(f"- Alerts: `outputs/alerts_{T}.json`")
    md.append(f"- Claim evidence: `outputs/claim_evidence_{T}.html`")
    md.append("")
    md.append("## 10) Next steps (what a human should do)")
    md.append("1) Open the **dashboard** first. Read rating + red flags.")
    md.append("2) Open the **news clickpack**. Click the top negative headlines and confirm they’re real + recent.")
    md.append("3) If your thesis depends on a specific risk (labor/regulatory/insurance), open **alerts** + **claim evidence**.")
    md.append("4) If anything looks off, treat score as directional and verify via earnings + filings.")
    md.append("")

    md_text = "\n".join(md)
    md_path = OUTPUTS / f"{T}_SUPER_Memo.md"
    md.append(_de_jargon(STORYTIME_WALKTHROUGH_BLOCK))

md_path.write_text(_de_jargon(_super_postprocess_text(md_text)), encoding="utf-8")

    doc = Document()
    doc.add_heading(f"SUPER Investment Memo — {T}", level=1)
    for line in md:
        if line.startswith("# "):
            continue
        if line.startswith("## "):
            doc.add_heading(line.replace("## ",""), level=2)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line.strip() == "":
            doc.add_paragraph("")
        else:
            doc.add_paragraph(line)

    doc_path = EXPORT / f"{T}_SUPER_Memo.docx"
    doc.save(doc_path)


    print('DONE ✅ SUPER memo created:')
    print(f"- {md_path}")
    print(f"- {doc_path}")

# (PDF export disabled — handled by soffice in wrapper)


if __name__ == "__main__":
    ap = argparse.ArgumentParser()
    ap.add_argument("--ticker", required=True)
    ap.add_argument("--thesis", required=True)
    args = ap.parse_args()
    main(args.ticker, args.thesis)

