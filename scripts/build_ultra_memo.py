#!/usr/bin/env python3

import json
import argparse
from pathlib import Path
from datetime import datetime
import pandas as pd
from docx import Document

ROOT = Path(__file__).resolve().parents[1]
OUTPUTS = ROOT / "outputs"
EXPORT = ROOT / "export"
DATA = ROOT / "data" / "processed"

def money(x):
    if x is None or (isinstance(x,float) and pd.isna(x)):
        return "N/A"
    return f"${x/1_000_000_000:,.2f}B"

def pct(x):
    if x is None or (isinstance(x,float) and pd.isna(x)):
        return "N/A"
    return f"{x:.2f}%"

def main(ticker, thesis_path):
    T = ticker.upper()

    comps = pd.read_csv(DATA / "comps_snapshot.csv")
    row = comps[comps["ticker"].str.upper()==T]
    if row.empty:
        raise ValueError("Ticker not found in comps_snapshot")
    row = row.iloc[0]

    rev = row.get("revenue_ttm_yoy_pct")
    fcf = row.get("fcf_ttm")
    margin = row.get("fcf_margin_ttm_pct")

    fy = row.get("fcf_yield_pct")
    if fy is None:
        raw = row.get("fcf_yield")
        try:
            fy = float(raw) * 100.0 if raw is not None else None
        except:
            fy = None

    thesis = json.load(open(thesis_path))
    thesis_text = thesis.get("description","N/A")

    now = datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC")

    md = f"""ULTRA Investment Memo — {T}
Generated: {now}

Your thesis
{thesis_text}

Core Numbers
Revenue growth YoY: {pct(rev)}
Free cash flow (TTM): {money(fcf)}
FCF margin: {pct(margin)}
FCF yield: {pct(fy)}

What this means (super simple)

Revenue growth = are people buying more?

Free cash flow = after paying bills, is cash left?

FCF margin = how much cash per $100 sales.

FCF yield = how cheap stock is vs cash.
"""

    md_path = OUTPUTS / f"{T}_ULTRA_Memo.md"
    md_path.write_text(md)

    doc = Document()
    doc.add_heading(f"ULTRA Investment Memo — {T}", level=1)
    doc.add_paragraph(md)
    doc_path = EXPORT / f"{T}_ULTRA_Memo.docx"
    doc.save(doc_path)

    print("DONE ✅ ULTRA memo created:")
    print(md_path)
    print(doc_path)

if __name__ == "__main__":
    ap = argparse.ArgumentParser()
    ap.add_argument("--ticker", required=True)
    ap.add_argument("--thesis", required=True)
    args = ap.parse_args()
    main(args.ticker, args.thesis)
