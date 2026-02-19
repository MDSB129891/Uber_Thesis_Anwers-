import argparse, json, math, subprocess
from pathlib import Path
from datetime import datetime, timezone

import pandas as pd
from docx import Document

def _pick_text(d: dict, keys):
    for k in keys:
        v = d.get(k)
        if isinstance(v, str) and v.strip():
            return v.strip()
    return ""



ROOT = Path(__file__).resolve().parents[1]

def money(x):
    if x is None or (isinstance(x, float) and math.isnan(x)):
        return "N/A"
    x = float(x)
    ax = abs(x)
    if ax >= 1e12: return f"${x/1e12:.2f}T"
    if ax >= 1e9:  return f"${x/1e9:.2f}B"
    if ax >= 1e6:  return f"${x/1e6:.2f}M"
    return f"${x:,.0f}"

def pct(x):
    if x is None or (isinstance(x, float) and math.isnan(x)):
        return "N/A"
    return f"{float(x):.2f}%"

def num(x):
    if x is None or (isinstance(x, float) and math.isnan(x)):
        return "N/A"
    return f"{float(x):.2f}"

def verdict_band(value, bands):
    """
    bands: list of tuples (name, predicate) in priority order.
    returns (label, emoji)
    """
    if value is None or (isinstance(value, float) and math.isnan(value)):
        return ("UNKNOWN", "❓")
    for label, emoji, pred in bands:
        if pred(value):
            return (label, emoji)
    return ("UNKNOWN", "❓")

def load_json(path: Path):
    if not path.exists():
        return {}
    return json.loads(path.read_text(encoding="utf-8"))

def load_comps_row(ticker: str):
    p = ROOT / "data/processed/comps_snapshot.csv"
    df = pd.read_csv(p)
    df["ticker"] = df["ticker"].astype(str).str.upper().str.strip()
    r = df[df["ticker"] == ticker.upper().strip()]
    if len(r) == 0:
        raise ValueError(f"Ticker {ticker} not found in {p}")
    return r.iloc[0].to_dict()

def build_md(ticker: str, thesis_text: str):
    T = ticker.upper().strip()
    now = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")

    comps = load_comps_row(T)

    decision = load_json(ROOT / "outputs/decision_summary.json")
    veracity = load_json(ROOT / "outputs" / f"veracity_{T}.json")
    risk = load_json(ROOT / "outputs" / f"news_risk_summary_{T}.json")

    # Core metrics
    rev_y = comps.get("revenue_ttm_yoy_pct", None)
    fcf_ttm = comps.get("fcf_ttm", None)
    fcf_margin = comps.get("fcf_margin_ttm_pct", None)

    # fcf_yield may be stored either as pct or decimal; prefer pct column if present
    fcf_yield_pct = comps.get("fcf_yield_pct", None)
    if fcf_yield_pct is None or (isinstance(fcf_yield_pct, float) and math.isnan(fcf_yield_pct)):
        fy = comps.get("fcf_yield", None)
        if fy is not None and not (isinstance(fy, float) and math.isnan(fy)):
            # assume decimal (0.06 -> 6%)
            fcf_yield_pct = float(fy) * 100.0

    mcap = comps.get("market_cap", None)
    cash = comps.get("cash", None)
    debt = comps.get("debt", None)
    net_debt = comps.get("net_debt", None)
    nd_to_fcf = comps.get("net_debt_to_fcf_ttm", None)

    rating = decision.get("rating", "N/A")
    score = decision.get("score", "N/A")
    vscore = veracity.get("confidence_score")
    if vscore is None:
        vscore = veracity.get("score")
    if vscore is None:
        vscore = veracity.get("veracity_score")
    if vscore is None:
        vscore = veracity.get("confidence")
    if vscore is None:
        vscore = veracity.get("evidence_confidence")
    if vscore is None:
        vscore = "N/A"


    news_shock_30d = risk.get("news_shock_30d", None)
    labor_30d = risk.get("risk_labor_neg_30d", None)
    reg_30d = risk.get("risk_regulatory_neg_30d", None)
    ins_30d = risk.get("risk_insurance_neg_30d", None)

    # Linked Good/Bad
    bands_rev = [
        ("GOOD", "✅", lambda v: v > 10),
        ("OK", "🟡", lambda v: 0 <= v <= 10),
        ("BAD", "❌", lambda v: v < 0),
    ]
    bands_fcf = [
        ("GOOD", "✅", lambda v: v > 0),
        ("BAD", "❌", lambda v: v < 0),
    ]
    bands_margin = [
        ("GOOD", "✅", lambda v: v >= 10),
        ("OK", "🟡", lambda v: 3 <= v < 10),
        ("BAD", "❌", lambda v: v <= 0),
    ]
    bands_yield = [
        ("CHEAP", "✅", lambda v: v > 5),
        ("NEUTRAL", "🟡", lambda v: 2 <= v <= 5),
        ("EXPENSIVE", "❌", lambda v: v < 2),
    ]
    bands_debt = [
        ("GOOD", "✅", lambda v: v < 3),
        ("WATCH", "🟡", lambda v: 3 <= v <= 6),
        ("HIGH RISK", "❌", lambda v: v > 6),
    ]
    # For headline shock: closer to 0 is calmer; very negative = more ugly headlines
    bands_shock = [
        ("CALM", "✅", lambda v: v >= -15),
        ("WATCH", "🟡", lambda v: -25 <= v < -15),
        ("UGLY", "❌", lambda v: v < -25),
    ]
    # For risk counts (30d): simple frequency bins
    bands_risk_ct = [
        ("LOW", "✅", lambda v: v <= 2),
        ("WATCH", "🟡", lambda v: 3 <= v <= 5),
        ("HIGH", "❌", lambda v: v >= 6),
    ]

    rev_label, rev_emoji = verdict_band(rev_y, bands_rev)
    fcf_label, fcf_emoji = verdict_band(fcf_ttm, bands_fcf)
    mar_label, mar_emoji = verdict_band(fcf_margin, bands_margin)
    yld_label, yld_emoji = verdict_band(fcf_yield_pct, bands_yield)
    deb_label, deb_emoji = verdict_band(nd_to_fcf, bands_debt)
    shk_label, shk_emoji = verdict_band(news_shock_30d, bands_shock)
    lab_label, lab_emoji = verdict_band(labor_30d, bands_risk_ct)
    reg_label, reg_emoji = verdict_band(reg_30d, bands_risk_ct)
    ins_label, ins_emoji = verdict_band(ins_30d, bands_risk_ct)

    md = []
    md.append(f"# SUPERPLUS Investment Memo — {T}")
    md.append(f"*Generated: {now}*")
    md.append("")
    md.append("## 1) Your thesis (what you believe)")
    md.append(f"**{thesis_text.strip() if thesis_text.strip() else 'N/A'}**")
    md.append("")
    md.append("## 2) What the model concluded (plain English)")
    md.append(f"- **Rating:** **{rating}** (score **{score}/100**)")
    md.append(f"- **Evidence confidence / veracity:** **{vscore}** (higher = more trustworthy coverage)")
    md.append("")
    md.append("## 3) The 30-second explanation (for total beginners)")
    md.append("Think of this like a **car dashboard**:")
    md.append("- The **score** is the overall attractiveness estimate.")
    md.append("- The **buckets** explain *why* the score happened.")
    md.append("- The **news/risk** items try to spot headline landmines.")
    md.append("- The **thesis test** checks whether the facts match the story you’re betting on.")
    md.append("")
    md.append("## Good vs Bad cheat-sheet (linked to this ticker)")
    md.append("Each line shows: **rule band → today’s value → verdict**.")
    md.append("")
    md.append("### Sales growth compared to last year (revenue growth)")
    md.append("- Rule band: Usually good **> +10%** | OK **0% to +10%** | Usually bad **< 0%**")
    md.append(f"- **{T} today:** **{pct(rev_y)}** → **{rev_label}** {rev_emoji}")
    md.append("")
    md.append("### Cash left over after all bills in the last 12 months (free cash flow)")
    md.append("- Rule band: Good **positive** | Bad **negative**")
    md.append(f"- **{T} today:** **{money(fcf_ttm)}** → **{fcf_label}** {fcf_emoji}")
    md.append("")
    md.append("### Cash efficiency of sales (free cash flow margin)")
    md.append("- Rule band: Usually good **≥ 10%** | OK **3% to 10%** | Bad **≤ 0%**")
    md.append(f"- **{T} today:** **{pct(fcf_margin)}** → **{mar_label}** {mar_emoji}")
    md.append("")
    md.append("### Cash return vs stock price (free cash flow yield)")
    md.append("- Rule band: Often cheap **> 5%** | Neutral **2% to 5%** | Often expensive **< 2%**")
    md.append(f"- **{T} today:** **{pct(fcf_yield_pct)}** → **{yld_label}** {yld_emoji}")
    md.append("")
    md.append("### Debt stress (net debt divided by free cash flow)")
    md.append("- Rule band: Good **< 3x** | Watch **3x to 6x** | High risk **> 6x**")
    md.append(f"- **{T} today:** **{num(nd_to_fcf)}x** → **{deb_label}** {deb_emoji}")
    md.append("")
    md.append("### Headline negativity in the last 30 days (news shock)")
    md.append("- Rule band: Calm **≥ -15** | Watch **-25 to -15** | Ugly **< -25**")
    md.append(f"- **{T} today:** **{num(news_shock_30d)}** → **{shk_label}** {shk_emoji}")
    md.append("")
    md.append("### Risk headline counts in the last 30 days")
    md.append("- Rule band: Low **0–2** | Watch **3–5** | High **6+**")
    md.append(f"- Labor risk headlines: **{labor_30d if labor_30d is not None else 'N/A'}** → **{lab_label}** {lab_emoji}")
    md.append(f"- Regulatory risk headlines: **{reg_30d if reg_30d is not None else 'N/A'}** → **{reg_label}** {reg_emoji}")
    md.append(f"- Insurance risk headlines: **{ins_30d if ins_30d is not None else 'N/A'}** → **{ins_label}** {ins_emoji}")
    md.append("")
    md.append("## 4) Core numbers (sanity-check)")
    md.append(f"- Sales growth compared to last year: **{pct(rev_y)}** _(comps_snapshot → revenue_ttm_yoy_pct)_")
    md.append(f"- Cash left over after all bills (last 12 months): **{money(fcf_ttm)}** _(comps_snapshot → fcf_ttm)_")
    md.append(f"- Cash efficiency of sales: **{pct(fcf_margin)}** _(comps_snapshot → fcf_margin_ttm_pct)_")
    md.append(f"- Cash return vs price paid: **{pct(fcf_yield_pct)}** _(comps_snapshot → fcf_yield)_")
    md.append("")
    md.append("## 5) Balance sheet snapshot (why debt matters)")
    md.append(f"- Market cap: **{money(mcap)}**")
    md.append(f"- Cash: **{money(cash)}**")
    md.append(f"- Debt: **{money(debt)}**")
    md.append(f"- Net debt (debt minus cash): **{money(net_debt)}**")
    md.append(f"- Net debt divided by free cash flow: **{num(nd_to_fcf)}x**")
    md.append("")
    md.append("## Storytime walkthrough (explain it like I’m five)")
    md.append(f"Okay. Imagine **{T}** is a **gigantic toy factory**.")
    md.append("You’re asking: *“Is this toy factory getting stronger… or about to hit expensive problems?”*")
    md.append("")
    md.append("### Step 1 — Are more toys being sold? (sales growth)")
    md.append(f"Today: **{pct(rev_y)}** → That tells us how sales changed compared to last year.")
    md.append("")
    md.append("### Step 2 — Is there money left in the piggy bank? (free cash flow)")
    md.append(f"Today: **{money(fcf_ttm)}** → After paying bills and investing, what’s left over.")
    md.append("")
    md.append("### Step 3 — Is the factory efficient? (free cash flow margin)")
    md.append(f"Today: **{pct(fcf_margin)}** → Out of every $100 of sales, how much becomes real cash.")
    md.append("")
    md.append("### Step 4 — Is the stock price cheap or expensive vs that cash? (free cash flow yield)")
    md.append(f"Today: **{pct(fcf_yield_pct)}** → Higher often means cheaper (but sometimes ‘cheap for a reason’).")
    md.append("")
    md.append("### Step 5 — Could debt cause stress if something goes wrong? (net debt / free cash flow)")
    md.append(f"Today: **{num(nd_to_fcf)}x** → Roughly how many years of current cash it would take to pay off net debt.")
    md.append("")
    md.append("## What to open (dopamine mode)")
    md.append(f"- Dashboard: `outputs/decision_dashboard_{T}.html`")
    md.append(f"- News clickpack: `outputs/news_clickpack_{T}.html`")
    md.append(f"- Claim evidence: `outputs/claim_evidence_{T}.html`")
    md.append("")
    return "\n".join(md)

def md_to_docx(md_text: str, docx_path: Path):
    doc = Document()
    for line in md_text.splitlines():
        line = line.rstrip()
        if not line:
            doc.add_paragraph("")
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        else:
            doc.add_paragraph(line)
    doc.save(str(docx_path))

def export_pdf(docx_path: Path, outdir: Path):
    soffice = Path("/opt/homebrew/bin/soffice")
    if not soffice.exists():
        return False
    try:
        subprocess.run([str(soffice), "--headless", "--convert-to", "pdf", "--outdir", str(outdir), str(docx_path)],
                       check=False, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
        return True
    except Exception:
        return False

def main(ticker: str, thesis_path: Path):
    T = ticker.upper().strip()
    thesis = json.loads(thesis_path.read_text(encoding="utf-8"))
    thesis_text = _pick_text(thesis, ("description","thesis","thesis_text","text","summary","narrative","prompt","name","title"))

    md = build_md(T, thesis_text)

    out_md = ROOT / "outputs" / f"{T}_SUPERPLUS_CLEAN_Memo.md"
    out_docx = ROOT / "export" / f"{T}_SUPERPLUS_CLEAN_Memo.docx"
    out_pdf = ROOT / "export" / f"{T}_SUPERPLUS_CLEAN_Memo.pdf"

    out_md.write_text(md, encoding="utf-8")
    md_to_docx(md, out_docx)
    export_pdf(out_docx, out_docx.parent)

    print("DONE ✅ SUPERPLUS CLEAN memo created:")
    print(f"- {out_md}")
    print(f"- {out_docx}")
    if out_pdf.exists():
        print(f"- {out_pdf}")
    else:
        print("⚠️ PDF not created (soffice issue). DOCX exists.")

if __name__ == "__main__":
    ap = argparse.ArgumentParser()
    ap.add_argument("--ticker", required=True)
    ap.add_argument("--thesis", required=True)
    args = ap.parse_args()
    main(args.ticker, Path(args.thesis))
