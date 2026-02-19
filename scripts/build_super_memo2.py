#!/usr/bin/env python3
import argparse, json, math, subprocess
from pathlib import Path
from datetime import datetime, timezone

import pandas as pd
from docx import Document

ROOT = Path(".")
OUTPUTS = ROOT / "outputs"
EXPORT = ROOT / "export"

def utc_now():
    return datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")

def _read_json(p: Path, default=None):
    try:
        return json.loads(p.read_text(encoding="utf-8"))
    except Exception:
        return default if default is not None else {}

def _fmt_money(x):
    if x is None or (isinstance(x, float) and math.isnan(x)):
        return "N/A"
    try:
        x = float(x)
    except Exception:
        return "N/A"
    sign = "-" if x < 0 else ""
    x = abs(x)
    if x >= 1e12: return f"{sign}${x/1e12:.2f}T"
    if x >= 1e9:  return f"{sign}${x/1e9:.2f}B"
    if x >= 1e6:  return f"{sign}${x/1e6:.2f}M"
    if x >= 1e3:  return f"{sign}${x/1e3:.2f}K"
    return f"{sign}${x:.0f}"

def _fmt_pct(x):
    if x is None or (isinstance(x, float) and math.isnan(x)):
        return "N/A"
    try:
        return f"{float(x):.2f}%"
    except Exception:
        return "N/A"

def _export_pdf(docx_path: Path) -> Path:
    pdf_path = docx_path.with_suffix(".pdf")
    soffice = "/opt/homebrew/bin/soffice"
    try:
        subprocess.run(
            [soffice, "--headless", "--convert-to", "pdf", "--outdir", str(docx_path.parent), str(docx_path)],
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE
        )
        return pdf_path
    except Exception:
        return pdf_path  # best-effort; caller can check exists

def _docx_from_md(md_text: str, docx_path: Path, title: str):
    doc = Document()
    doc.add_heading(title, level=0)
    doc.add_paragraph(f"Generated: {utc_now()}")

    for line in md_text.splitlines():
        line = line.rstrip()
        if not line:
            doc.add_paragraph("")
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
            continue
        if line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
            continue
        doc.add_paragraph(line)

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))

def main(ticker: str, thesis_path: Path):
    T = ticker.upper()
    EXPORT.mkdir(parents=True, exist_ok=True)
    OUTPUTS.mkdir(parents=True, exist_ok=True)

    # Inputs
    thesis = _read_json(thesis_path, {})
    thesis_headline = thesis.get("name") or f"{T}: thesis"
    thesis_text = thesis.get("description") or thesis.get("thesis") or thesis.get("text") or thesis.get("headline") or "N/A"

    decision = _read_json(OUTPUTS / "decision_summary.json", {})
    score = decision.get("score")
    rating = decision.get("rating")
    bucket_scores = decision.get("bucket_scores", {}) or {}
    red_flags = decision.get("red_flags", []) or []

    veracity = _read_json(OUTPUTS / f"veracity_{T}.json", {})
    confidence = veracity.get("confidence", veracity.get("score", veracity.get("veracity_score")))

    # Comps snapshot core
    snap_path = ROOT / "data/processed/comps_snapshot.csv"
    snap_row = {}
    if snap_path.exists():
        df = pd.read_csv(snap_path)
        r = df[df["ticker"].astype(str).str.upper() == T]
        if len(r):
            snap_row = r.iloc[0].to_dict()

    revenue_growth = snap_row.get("revenue_ttm_yoy_pct")
    fcf_ttm = snap_row.get("fcf_ttm")
    fcf_margin = snap_row.get("fcf_margin_ttm_pct")
    fcf_yield = snap_row.get("fcf_yield_pct")
    if fcf_yield is None:
        fy = snap_row.get("fcf_yield")
        if fy is not None and not (isinstance(fy, float) and math.isnan(fy)):
            try:
                fcf_yield = float(fy) * 100.0
            except Exception:
                fcf_yield = None

    mktcap = snap_row.get("market_cap")
    cash = snap_row.get("cash")
    debt = snap_row.get("debt")
    net_debt = snap_row.get("net_debt")
    net_debt_to_fcf = snap_row.get("net_debt_to_fcf_ttm")

    # Claim evidence (optional)
    claim_pack = _read_json(OUTPUTS / f"claim_evidence_{T}.json", {})
    claim_results = claim_pack.get("results", []) if isinstance(claim_pack, dict) else []

    def good_bad_block():
        return """## Good vs Bad cheat-sheet (how to judge the numbers)

Think of every metric like a **warning light** on a car.

### 1) Revenue growth (sales compared to last year)
- ✅ Usually good: **above +10%**
- 🟡 Mixed: **0% to +10%**
- ❌ Usually bad: **below 0%**

### 2) Free cash flow (real leftover cash)
- ✅ Good: **positive and stable/increasing**
- 🟡 Mixed: **positive but volatile**
- ❌ Bad: **negative consistently**

### 3) Free cash flow margin (cash per $100 of sales)
- ✅ Good: **10% or higher** (industry-dependent)
- 🟡 Mixed: **3% to 10%**
- ❌ Bad: **0% or negative**

### 4) Free cash flow yield (cash return compared to stock price)
- ✅ Often “cheap”: **above 5%**
- 🟡 Neutral: **2% to 5%**
- ❌ Often “expensive”: **below 2%**

### 5) Net debt (debt minus cash)
- ✅ Better: low net debt (or net cash)
- ❌ Risk: big net debt + weakening cash flow

### 6) Net debt divided by free cash flow (years to pay debt)
- ✅ Good: **below 3x**
- 🟡 Watch: **3x to 6x**
- ❌ High risk: **above 6x**
"""

    def storytime_block():
        return f"""## Storytime walkthrough (explain it like I’m five)

Imagine **{T}** is a giant factory.

You’re asking: **“Will this factory be stronger later, or will costs/problems crush it?”**

### Step 1 — Are sales growing?
“Sales growth compared to last year” tells us if more people are buying the product.
- If it’s growing: demand is usually stronger.
- If it’s shrinking: demand may be weakening.

### Step 2 — Is there real leftover cash?
“Free cash flow” is the money left after paying bills **and** investing to keep the business running.
- Positive = the piggy bank is filling.
- Negative = the piggy bank is leaking.

### Step 3 — Is the factory efficient?
“Free cash flow margin” asks: out of every $100 of sales, how many dollars become leftover cash?

### Step 4 — Are we paying a fair price?
“Free cash flow yield” asks: how much real cash do you get compared to what you pay for the stock?

### Step 5 — Can debt become a problem?
“Net debt” and “years to pay debt” tell us how stressed the company could get in a downturn.

### Step 6 — News landmines
If headlines are very negative (labor, regulation, insurance), stocks can drop fast even if the business is okay.
"""

    md = []
    md.append(f"# SUPER Investment Memo — {T}")
    md.append(f"*Generated: {utc_now()}*")
    md.append("")
    md.append("## 1) Your thesis (what you believe)")
    md.append(f"**{thesis_headline}**")
    md.append(thesis_text)
    md.append("")
    md.append("## 2) What the model concluded (plain English)")
    md.append(f"- **Rating:** **{rating if rating is not None else 'N/A'}** (score **{score if score is not None else 'N/A'}/100**)")
    md.append(f"- **Evidence confidence:** **{confidence if confidence is not None else 'N/A'}** (higher = more trustworthy coverage)")
    md.append("")
    md.append("## 3) The 30-second explanation (for total beginners)")
    md.append("- The **score** is the overall “health + attractiveness” estimate.")
    md.append("- The **buckets** show *why* the score happened (cash, valuation, growth, quality, balance-sheet risk).")
    md.append("- The **news/risk** section tries to spot headline landmines.")
    md.append("- The **thesis test** checks whether the facts match the story you’re betting on.")
    md.append("")
    md.append("## 4) Core numbers (sanity-check)")
    md.append(f"- Sales growth compared to last year: **{_fmt_pct(revenue_growth)}** (source: comps_snapshot → revenue_ttm_yoy_pct)")
    md.append(f"- Free cash flow in the last twelve months: **{_fmt_money(fcf_ttm)}** (source: comps_snapshot → fcf_ttm)")
    md.append(f"- Free cash flow margin: **{_fmt_pct(fcf_margin)}** (source: comps_snapshot → fcf_margin_ttm_pct)")
    md.append(f"- Free cash flow yield: **{_fmt_pct(fcf_yield)}** (source: comps_snapshot → fcf_yield_pct / fcf_yield)")
    md.append("")
    md.append("## 5) Balance sheet snapshot (why debt matters)")
    md.append(f"- Market value (market capitalization): **{_fmt_money(mktcap)}**")
    md.append(f"- Cash: **{_fmt_money(cash)}**")
    md.append(f"- Debt: **{_fmt_money(debt)}**")
    md.append(f"- Net debt (debt minus cash): **{_fmt_money(net_debt)}**")
    md.append(f"- Net debt divided by free cash flow: **{(f'{float(net_debt_to_fcf):.2f}x' if net_debt_to_fcf is not None and not (isinstance(net_debt_to_fcf,float) and math.isnan(net_debt_to_fcf)) else 'N/A')}**")
    md.append("")
    md.append("## 6) Bucket scores (what drove the rating)")
    for k in ["cash_level","valuation","growth","quality","balance_risk"]:
        v = bucket_scores.get(k)
        if v is None:
            continue
        pretty = {
            "cash_level":"Cash strength",
            "valuation":"Valuation",
            "growth":"Growth",
            "quality":"Business quality",
            "balance_risk":"Balance-sheet risk"
        }.get(k, k)
        md.append(f"- **{pretty}:** **{v}**")
    md.append("")
    md.append("## 7) Red flags (things that can hurt stock fast)")
    if red_flags:
        for rf in red_flags:
            md.append(f"- {rf}")
    else:
        md.append("- None flagged by the engine.")
    md.append("")
    md.append(good_bad_block())
    md.append("")
    md.append(storytime_block())
    md.append("")

    md.append("## 8) Thesis test (PASS/FAIL vs your claims)")
    if claim_results:
        for r in claim_results[:12]:
            status = r.get("status","UNKNOWN")
            statement = r.get("statement","(no statement)")
            metric = r.get("metric","(no metric)")
            val = r.get("value")
            md.append(f"- **{status}** — {statement}  \n  Metric `{metric}` | Actual: **{val}**")
    else:
        md.append("- No claim results found (claim evidence pack missing or empty).")
    md.append("")

    md.append("## 9) What to open (dopamine mode)")
    md.append(f"- Dashboard: `outputs/decision_dashboard_{T}.html`")
    md.append(f"- News clickpack: `outputs/news_clickpack_{T}.html`")
    md.append(f"- Claim evidence: `outputs/claim_evidence_{T}.html`")
    md.append(f"- SUPER PDF: `export/{T}_SUPER_Memo2.pdf`")
    md.append("")
    md.append("## 10) Next steps (human checklist)")
    md.append("1) Open the dashboard. Look at rating + red flags.")
    md.append("2) Read this SUPER memo PDF. Understand what the numbers mean.")
    md.append("3) Open the news clickpack. Click the top negatives and confirm they’re real + recent.")
    md.append("4) If the thesis depends on labor/regulation/insurance, focus on those headlines first.")
    md.append("")

    md_text = "\n".join(md)

    md_path = OUTPUTS / f"{T}_SUPER_Memo2.md"
    docx_path = EXPORT / f"{T}_SUPER_Memo2.docx"
    pdf_path = EXPORT / f"{T}_SUPER_Memo2.pdf"

    md_path.write_text(md_text, encoding="utf-8")
    _docx_from_md(md_text, docx_path, f"SUPER Investment Memo — {T}")
    out_pdf = _export_pdf(docx_path)
    if out_pdf.exists():
        pdf_path = out_pdf

    print("DONE ✅ SUPER memo2 created:")
    print(f"- {md_path}")
    print(f"- {docx_path}")
    if pdf_path.exists():
        print(f"- {pdf_path}")
    else:
        print("⚠️ PDF not found after export (LibreOffice conversion may have failed).")

if __name__ == "__main__":
    ap = argparse.ArgumentParser()
    ap.add_argument("--ticker", required=True)
    ap.add_argument("--thesis", required=True)
    args = ap.parse_args()
    main(args.ticker, Path(args.thesis))
