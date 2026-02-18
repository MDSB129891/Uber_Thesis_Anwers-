#!/usr/bin/env python3
from __future__ import annotations

import json
import re
import os
from pathlib import Path
from urllib.parse import urlparse
from datetime import datetime
from typing import Dict, Tuple, List, Optional

import pandas as pd

# Word + Excel
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from openpyxl import Workbook
from openpyxl.utils.dataframe import dataframe_to_rows
from openpyxl.styles import Font, Alignment
from openpyxl.worksheet.table import Table, TableStyleInfo


ROOT = Path(__file__).resolve().parents[1]
DATA_PROCESSED = ROOT / "data" / "processed"
OUTPUTS = ROOT / "outputs"
EXPORT = ROOT / "export"

# You can override the ticker for reports:
#   TICKER=LYFT python3 scripts/build_investment_report.py
TICKER = os.getenv("TICKER", "UBER").upper()

WHITELIST_PATH = EXPORT / "source_whitelist.csv"   # editable by you (no code)
DEFAULT_TOP_TIER_DOMAINS = {
    "reuters.com",
    "bloomberg.com",
    "wsj.com",
    "ft.com",
    "economist.com",
    "cnbc.com",
    "nytimes.com",
    "washingtonpost.com",
    "theinformation.com",
    "techcrunch.com",
    "theverge.com",
    "forbes.com",
    "marketwatch.com",
    "sec.gov",
}

REDIRECT_LIKE_PATTERNS = [
    r"finnhub\.io/api/news\?id=",
    r"^None$",
    r"^$",
]

_WS = re.compile(r"\s+")
_PUNCT = re.compile(r"[^\w\s]")


# ---------------------------
# Basic utilities
# ---------------------------
def ensure_dirs():
    OUTPUTS.mkdir(parents=True, exist_ok=True)
    EXPORT.mkdir(parents=True, exist_ok=True)


def load_json(path: Path) -> dict:
    if not path.exists():
        return {}
    return json.loads(path.read_text(encoding="utf-8"))


def safe_read_csv(path: Path) -> pd.DataFrame:
    if not path.exists():
        return pd.DataFrame()
    return pd.read_csv(path)


def normalize_title(title: str) -> str:
    t = (title or "").strip().lower()
    t = _PUNCT.sub(" ", t)
    t = _WS.sub(" ", t).strip()
    return t


def extract_domain(url: str) -> str:
    if not url or str(url).strip() in ("None", ""):
        return ""
    try:
        u = str(url).strip()
        if not u.startswith("http"):
            return ""
        netloc = urlparse(u).netloc.lower()
        if netloc.startswith("www."):
            netloc = netloc[4:]
        return netloc
    except Exception:
        return ""


def is_redirect_like(url: str) -> bool:
    u = (url or "").strip()
    for pat in REDIRECT_LIKE_PATTERNS:
        if re.search(pat, u):
            return True
    return False


# ---------------------------
# Source whitelist (editable, no-code)
# ---------------------------
def load_source_whitelist() -> pd.DataFrame:
    """
    Editable by you at: export/source_whitelist.csv

    Columns:
      - domain
      - tier  (TOP / MID / LOW)
      - weight (numeric; optional)
    """
    WHITELIST_PATH.parent.mkdir(parents=True, exist_ok=True)

    if WHITELIST_PATH.exists():
        df = pd.read_csv(WHITELIST_PATH)
        df["domain"] = df["domain"].astype(str).str.lower().str.strip()
        df["tier"] = df.get("tier", "TOP").astype(str).str.upper().str.strip()
        df["weight"] = pd.to_numeric(df.get("weight", 1.0), errors="coerce").fillna(1.0)
        df = df[df["domain"] != ""].copy()
        return df

    starter = pd.DataFrame(
        [{"domain": d, "tier": "TOP", "weight": 1.0} for d in sorted(DEFAULT_TOP_TIER_DOMAINS)]
    )
    starter.to_csv(WHITELIST_PATH, index=False)
    return starter


def domain_tier(domain: str, whitelist: pd.DataFrame) -> str:
    if not domain:
        return "UNKNOWN"
    if whitelist is None or whitelist.empty:
        return "UNKNOWN"
    m = whitelist[whitelist["domain"].apply(lambda d: domain == d or domain.endswith("." + d))]
    if m.empty:
        return "OTHER"
    tier_order = {"TOP": 3, "MID": 2, "LOW": 1}
    m = m.copy()
    m["_rank"] = m["tier"].map(tier_order).fillna(0)
    best = m.sort_values("_rank", ascending=False).iloc[0]
    return str(best["tier"]).upper()


def is_top_tier(domain: str, whitelist: pd.DataFrame) -> bool:
    return domain_tier(domain, whitelist) == "TOP"


# ---------------------------
# Veracity metrics + Confidence Score
# ---------------------------
def build_veracity_summary(news_df: pd.DataFrame, whitelist: pd.DataFrame) -> dict:
    if news_df.empty:
        return {
            "total_rows": 0,
            "source_counts": {},
            "domain_counts_top10": {},
            "top_tier_rows": 0,
            "top_tier_pct": 0.0,
            "redirect_like_rows": 0,
            "missing_url_rows": 0,
            "duplicate_titles_est": 0,
        }

    df = news_df.copy()
    df["source"] = df["source"].astype(str).fillna("")
    df["url"] = df["url"].astype(str).fillna("")
    df["domain"] = df["url"].apply(extract_domain)
    df["is_top_tier"] = df["domain"].apply(lambda d: is_top_tier(d, whitelist))
    df["is_redirect_like"] = df["url"].apply(is_redirect_like)
    df["missing_url"] = df["url"].apply(lambda x: (not x) or str(x).strip() in ("None", ""))

    total = len(df)
    top_tier_rows = int(df["is_top_tier"].sum())
    redirect_like_rows = int(df["is_redirect_like"].sum())
    missing_url_rows = int(df["missing_url"].sum())

    source_counts = df["source"].value_counts().to_dict()
    domain_counts = df["domain"].value_counts().head(10).to_dict()

    df["title_norm"] = df["title"].astype(str).apply(normalize_title)
    dup_est = int(df.duplicated(subset=["ticker", "title_norm"]).sum())

    return {
        "total_rows": total,
        "source_counts": source_counts,
        "domain_counts_top10": domain_counts,
        "top_tier_rows": top_tier_rows,
        "top_tier_pct": round((top_tier_rows / total) * 100.0, 2) if total else 0.0,
        "redirect_like_rows": redirect_like_rows,
        "missing_url_rows": missing_url_rows,
        "duplicate_titles_est": dup_est,
    }


def compute_confidence_score(veracity: dict) -> Tuple[int, List[str]]:
    reasons = []
    total = int(veracity.get("total_rows", 0))
    top_pct = float(veracity.get("top_tier_pct", 0.0))
    redirect = int(veracity.get("redirect_like_rows", 0))
    missing = int(veracity.get("missing_url_rows", 0))
    dup = int(veracity.get("duplicate_titles_est", 0))

    score = 70

    if total >= 50 and top_pct >= 25:
        score += 10
        reasons.append("Good volume + decent share of top-tier sources.")
    elif total >= 50 and top_pct < 10:
        score -= 10
        reasons.append("Most news is not top-tier (consider editing whitelist tiers).")
    elif total < 15:
        score -= 10
        reasons.append("Low news volume (signal may be incomplete).")

    if redirect > 0:
        score -= min(10, redirect // 10 + 2)
        reasons.append("Some URLs are redirect-like; verification is harder.")
    if missing > 0:
        score -= min(10, missing // 10 + 2)
        reasons.append("Some items have missing URLs.")

    if dup > 0:
        score -= min(10, dup // 20 + 2)
        reasons.append("Duplicate headlines detected (syndication noise).")

    score = max(0, min(100, int(round(score))))
    if not reasons:
        reasons = ["Good coverage and verifiability."]
    return score, reasons


# ---------------------------
# Evidence tables + Curated pack
# ---------------------------
def build_evidence_tables(
    news_df: pd.DataFrame, ticker: str, whitelist: pd.DataFrame, days: int = 30, max_rows: int = 300
) -> Tuple[pd.DataFrame, pd.DataFrame]:
    if news_df.empty:
        cols = ["published_at", "ticker", "source", "tier", "risk_tag", "impact_score", "title", "url", "domain"]
        empty = pd.DataFrame(columns=cols)
        return empty, empty

    df = news_df.copy()
    df["ticker"] = df["ticker"].astype(str).str.upper()
    df = df[df["ticker"] == ticker.upper()].copy()

    df["published_at"] = pd.to_datetime(df["published_at"], errors="coerce", utc=True)
    cutoff = pd.Timestamp.utcnow() - pd.Timedelta(days=days)
    df = df[df["published_at"] >= cutoff].copy()

    df["impact_score"] = pd.to_numeric(df.get("impact_score", 0), errors="coerce").fillna(0).astype(int)
    df["risk_tag"] = df.get("risk_tag", "OTHER").astype(str).fillna("OTHER")
    df["source"] = df.get("source", "unknown").astype(str).fillna("unknown")
    df["title"] = df.get("title", "").astype(str).fillna("")
    df["url"] = df.get("url", "").astype(str).fillna("")
    df["domain"] = df["url"].apply(extract_domain)
    df["tier"] = df["domain"].apply(lambda d: domain_tier(d, whitelist))

    df = df.sort_values(["published_at"], ascending=[False])

    keep_cols = ["published_at", "ticker", "source", "tier", "risk_tag", "impact_score", "title", "url", "domain"]
    all_df = df[keep_cols].head(max_rows).copy()
    top_df = df[df["tier"] == "TOP"][keep_cols].head(max_rows).copy()

    for d in (all_df, top_df):
        d["published_at"] = pd.to_datetime(d["published_at"], errors="coerce", utc=True).dt.strftime("%Y-%m-%d %H:%M UTC")

    return top_df, all_df


def curate_evidence(evidence_df: pd.DataFrame, top_n: int = 12) -> pd.DataFrame:
    """
    Curate a small set of items so validation is painless.
    Prioritize: severe negatives + important risk tags + recency.
    """
    if evidence_df is None or evidence_df.empty:
        return pd.DataFrame(columns=["published_at", "risk_tag", "impact_score", "source", "tier", "title", "url"])

    df = evidence_df.copy()
    df["impact_score"] = pd.to_numeric(df.get("impact_score", 0), errors="coerce").fillna(0).astype(int)
    df["risk_tag"] = df.get("risk_tag", "OTHER").astype(str).str.upper()
    df["tier"] = df.get("tier", "").astype(str).str.upper()

    tag_weight = {"REGULATORY": 3, "INSURANCE": 3, "LABOR": 2, "SAFETY": 2, "OTHER": 1}
    df["_tag_w"] = df["risk_tag"].map(tag_weight).fillna(1)

    # recency proxy from published_at string (YYYY-mm-dd...)
    try:
        df["_dt"] = pd.to_datetime(df["published_at"], errors="coerce")
        df["_recency"] = df["_dt"].rank(ascending=False, method="dense")
    except Exception:
        df["_recency"] = 999

    # Priority: more negative impact is more urgent to verify.
    # If your impact_score uses negative for negatives, this works well.
    # We also bias slightly toward TOP tier when present.
    df["_tier_w"] = df["tier"].map({"TOP": 2, "MID": 1, "LOW": 0}).fillna(0)

    df["_priority"] = (-df["impact_score"]) * 4 + df["_tag_w"] * 3 + df["_tier_w"] * 2 + (1000 - df["_recency"])
    df = df.sort_values("_priority", ascending=False)

    keep = ["published_at", "risk_tag", "impact_score", "source", "tier", "title", "url"]
    return df[keep].head(top_n).copy()


# ---------------------------
# Bucket evidence + Metric Cheat Sheet
# ---------------------------
def build_bucket_evidence(decision_summary: dict) -> dict:
    buckets = decision_summary.get("bucket_scores", {}) or {}
    red_flags = decision_summary.get("red_flags", []) or []
    news_summary = decision_summary.get("news_summary", {}) or {}
    top_neg = news_summary.get("top_negative_titles_7d", []) or []

    drivers = {
        "cash_level": [
            "Driven by TTM free cash flow level (how much real cash the business produces).",
            "Higher, stable FCF usually increases score.",
        ],
        "valuation": [
            "Driven by FCF yield and peer ranking (cheaper vs cash = better).",
            "If price rises faster than cash, valuation bucket weakens.",
        ],
        "growth": [
            "Driven by TTM revenue YoY and FCF YoY + peer ranks.",
            "Slowing growth reduces score even if the company is profitable.",
        ],
        "quality": [
            "Driven by TTM FCF margin + peer rank.",
            "Higher margin = better business economics.",
        ],
        "balance_risk": [
            "Debt load vs TTM FCF (net-debt-to-FCF proxy).",
            "Negative news frequency/shock over last 7 days.",
            "Repeating risk tags (REGULATORY / INSURANCE / LABOR).",
            "Headline sentiment proxy score (7d/30d).",
        ],
    }

    return {
        "bucket_scores": buckets,
        "red_flags": red_flags,
        "drivers": drivers,
        "top_negative_headlines_7d": top_neg[:10],
    }


def metric_cheat_sheet() -> pd.DataFrame:
    """
    Beginner-friendly “good/ok/bad” key.
    These are rules of thumb, not universal truths.
    """
    rows = [
        {
            "Metric": "Revenue Growth (YoY %)",
            "Meaning (plain English)": "Is the company getting bigger?",
            "Good": "> 10%",
            "OK": "0–10%",
            "Bad": "< 0%",
            "Why it matters": "Growth supports future cash and valuation."
        },
        {
            "Metric": "Free Cash Flow (FCF)",
            "Meaning (plain English)": "Cash left after running the business and capex.",
            "Good": "Positive and rising",
            "OK": "Positive but flat",
            "Bad": "Negative",
            "Why it matters": "FCF funds buybacks, debt paydown, and growth."
        },
        {
            "Metric": "FCF Margin (%)",
            "Meaning (plain English)": "How much cash is produced per $1 of sales.",
            "Good": "> 10%",
            "OK": "0–10%",
            "Bad": "< 0%",
            "Why it matters": "Higher = better business economics."
        },
        {
            "Metric": "FCF Yield",
            "Meaning (plain English)": "Cash return vs price you pay (FCF / Market Cap).",
            "Good": "> 5%",
            "OK": "2–5%",
            "Bad": "< 2%",
            "Why it matters": "Higher yield often means “cheaper” vs cash."
        },
        {
            "Metric": "Net Debt / FCF",
            "Meaning (plain English)": "How many years of FCF to pay off net debt.",
            "Good": "< 3x",
            "OK": "3–6x",
            "Bad": "> 6x",
            "Why it matters": "Lower = safer in downturns."
        },
        {
            "Metric": "News Neg Count (7d/30d)",
            "Meaning (plain English)": "How often negative stories show up.",
            "Good": "Low / declining",
            "OK": "Stable",
            "Bad": "Rising fast",
            "Why it matters": "Repetition suggests persistent risk."
        },
        {
            "Metric": "News Shock (7d/30d)",
            "Meaning (plain English)": "How severe/negative the headlines are.",
            "Good": "Near 0",
            "OK": "-1 to -10",
            "Bad": "< -10",
            "Why it matters": "Large negative shocks can signal real events."
        },
    ]
    return pd.DataFrame(rows)


# ---------------------------
# Pretty Word helpers
# ---------------------------
def _shade_cell(cell, fill_hex: str = "F2F2F2"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _format_table(table, header_fill="E8EEF7", header_font_size=10, body_font_size=9):
    hdr = table.rows[0].cells
    for c in hdr:
        _shade_cell(c, header_fill)
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(header_font_size)

    for row in table.rows[1:]:
        for c in row.cells:
            for p in c.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                for r in p.runs:
                    r.font.size = Pt(body_font_size)


# ---------------------------
# Human units + WARNING FIX INCLUDED
# ---------------------------
def add_human_units_comps(comps: pd.DataFrame) -> pd.DataFrame:
    if comps is None or comps.empty:
        return comps

    df = comps.copy()

    # numeric coercion (safe) — no FutureWarning
    for c in df.columns:
        if c in ("ticker", "period_end"):
            continue
        try:
            df[c] = pd.to_numeric(df[c], errors="raise")
        except Exception:
            pass

    usd_cols = ["market_cap", "revenue_ttm", "fcf_ttm", "cash", "debt", "net_debt"]
    for c in usd_cols:
        if c in df.columns:
            df[f"{c}_usd_bn"] = pd.to_numeric(df[c], errors="coerce") / 1e9

    return df


# ---------------------------
# HTML report (clickable)
# ---------------------------
def write_html_report(
    out_path: Path,
    title: str,
    decision_summary: dict,
    comps: pd.DataFrame,
    risk_dash: pd.DataFrame,
    veracity: dict,
    confidence_score: int,
    confidence_reasons: List[str],
    curated_pack: pd.DataFrame,
    evidence_top: pd.DataFrame,
    evidence_all: pd.DataFrame,
):
    out_path.parent.mkdir(parents=True, exist_ok=True)

    score = decision_summary.get("score")
    rating = decision_summary.get("rating")
    buckets = decision_summary.get("bucket_scores", {})
    red_flags = decision_summary.get("red_flags", [])

    def df_to_html(df: pd.DataFrame, max_rows=50, escape=True):
        if df is None or df.empty:
            return "<p><em>No data.</em></p>"
        return df.head(max_rows).to_html(index=False, escape=escape)

    def evidence_to_html(df: pd.DataFrame, max_rows=100):
        if df is None or df.empty:
            return "<p><em>No evidence rows.</em></p>"
        d = df.head(max_rows).copy()

        def linkify(row):
            url = row.get("url", "")
            title = row.get("title", "")
            if url and url != "None" and str(url).startswith("http"):
                return f'<a href="{url}" target="_blank" rel="noopener noreferrer">{title}</a>'
            return title

        d["title"] = d.apply(linkify, axis=1)
        return d.to_html(index=False, escape=False)

    html = f"""
    <html>
      <head>
        <meta charset="utf-8"/>
        <title>{title}</title>
        <style>
          body {{ font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Arial, sans-serif; padding: 24px; line-height: 1.35; }}
          h1 {{ margin: 0 0 6px 0; }}
          h2 {{ margin-top: 28px; }}
          .pill {{ display: inline-block; padding: 6px 10px; border-radius: 999px; background: #f3f3f3; margin-right: 8px; font-weight: 600; }}
          .score {{ font-size: 28px; font-weight: 800; }}
          .muted {{ color: #666; }}
          table {{ border-collapse: collapse; width: 100%; }}
          th, td {{ border: 1px solid #ddd; padding: 8px; font-size: 13px; vertical-align: top; }}
          th {{ background: #f5f5f5; }}
          tr:nth-child(even) {{ background: #fafafa; }}
          .grid {{ display: grid; grid-template-columns: 1fr 1fr; gap: 16px; }}
          .card {{ border: 1px solid #e6e6e6; border-radius: 12px; padding: 14px; background: white; }}
          pre {{ white-space: pre-wrap; }}
        </style>
      </head>
      <body>
        <h1>{title}</h1>
        <div class="muted">Generated: {datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC")}</div>

        <h2>Decision</h2>
        <div class="grid">
          <div class="card">
            <div class="score">{score} / 100</div>
            <div><span class="pill">{rating}</span><span class="muted">Ticker: {TICKER}</span></div>
            <p><b>Buckets:</b> {buckets}</p>
            <p><b>Red flags:</b> {red_flags if red_flags else "None"}</p>
          </div>
          <div class="card">
            <h3 style="margin-top:0;">Confidence (Veracity)</h3>
            <div class="score">{confidence_score} / 100</div>
            <ul>
              {''.join([f"<li>{r}</li>" for r in confidence_reasons])}
            </ul>
            <p class="muted">This measures how verifiable the evidence is (not a price prediction).</p>
          </div>
        </div>

        <h2>Curated Evidence Pack (start here)</h2>
        <p class="muted">These are the highest-priority items to verify first. Click a headline to open it.</p>
        {evidence_to_html(curated_pack, max_rows=50)}

        <h2>Peer Comps</h2>
        {df_to_html(comps, max_rows=12)}

        <h2>Risk Dashboard</h2>
        {df_to_html(risk_dash, max_rows=100)}

        <h2>Veracity / Source Quality</h2>
        <div class="card">
          <pre>{json.dumps(veracity, indent=2)}</pre>
          <p class="muted">Edit whitelist: export/source_whitelist.csv (TOP/MID/LOW).</p>
        </div>

        <h2>Evidence (TOP tier)</h2>
        {evidence_to_html(evidence_top, max_rows=120)}

        <h2>Evidence (All)</h2>
        {evidence_to_html(evidence_all, max_rows=200)}

      </body>
    </html>
    """
    out_path.write_text(html, encoding="utf-8")


# ---------------------------
# WORD report (pretty + detailed)
# ---------------------------
def write_word_report(
    out_path: Path,
    decision_summary: dict,
    comps: pd.DataFrame,
    proxy: pd.DataFrame,
    risk_dash: pd.DataFrame,
    veracity: dict,
    confidence_score: int,
    confidence_reasons: List[str],
    bucket_evidence: dict,
    curated_pack: pd.DataFrame,
    evidence_top: pd.DataFrame,
):
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    doc.add_heading(f"Investment Decision Report — {TICKER}", level=0)
    doc.add_paragraph(f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}")

    score = decision_summary.get("score")
    rating = decision_summary.get("rating")
    buckets = decision_summary.get("bucket_scores", {}) or {}
    red_flags = decision_summary.get("red_flags", []) or {}
    news_summary = decision_summary.get("news_summary", {}) or {}

    # Pull key numbers from comps (beginner-friendly)
    key_nums = {}
    try:
        c = comps[comps["ticker"] == TICKER].iloc[0].to_dict()

        def pick(*names):
            for n in names:
                if n in c and pd.notna(c[n]):
                    return c[n]
            return None

        key_nums = {
            "Revenue (TTM)": pick("revenue_ttm_usd_bn", "revenue_ttm"),
            "Free Cash Flow (TTM)": pick("fcf_ttm_usd_bn", "fcf_ttm"),
            "FCF Margin (TTM %)": pick("fcf_margin_ttm_pct"),
            "Revenue Growth (TTM YoY %)": pick("revenue_ttm_yoy_pct"),
            "FCF Growth (TTM YoY %)": pick("fcf_ttm_yoy_pct"),
            "Market Cap": pick("market_cap_usd_bn", "market_cap"),
            "FCF Yield": pick("fcf_yield"),
            "Net Debt": pick("net_debt_usd_bn", "net_debt"),
            "Net Debt / FCF": pick("net_debt_to_fcf_ttm"),
        }
    except Exception:
        pass

    def fmt_money_bn(x):
        if x is None or (isinstance(x, float) and pd.isna(x)):
            return "N/A"
        try:
            x = float(x)
            if x > 1e6:
                return f"${x/1e9:.2f}B"
            return f"${x:.2f}B"
        except Exception:
            return str(x)

    def fmt_pct(x):
        if x is None or (isinstance(x, float) and pd.isna(x)):
            return "N/A"
        try:
            return f"{float(x):.2f}%"
        except Exception:
            return str(x)

    def fmt_ratio(x):
        if x is None or (isinstance(x, float) and pd.isna(x)):
            return "N/A"
        try:
            return f"{float(x):.2f}"
        except Exception:
            return str(x)

    # ---------------------------
    # Executive Summary
    # ---------------------------
    doc.add_heading("Executive Summary (plain English)", level=1)
    doc.add_paragraph(
        f"Today, {TICKER} scores {score}/100 with a rating of {rating}. "
        "This score combines five areas into one view: cash generation, valuation, growth, business quality, and risk."
    )

    takeaways = []
    if score is not None:
        if score >= 80:
            takeaways.append("Overall signal is strong: the numbers support the thesis today.")
        elif score >= 65:
            takeaways.append("Overall signal is mixed: some positives, some caution signs.")
        else:
            takeaways.append("Overall signal is weak: risk/price mismatch or weak fundamentals.")

    if red_flags:
        takeaways.append("Red flags were detected; review them before acting.")
    else:
        takeaways.append("No major red flags were detected by the model today.")

    for t in takeaways:
        doc.add_paragraph(t, style="List Bullet")

    if red_flags:
        for rf in red_flags:
            doc.add_paragraph(str(rf), style="List Bullet 2")

    # ---------------------------
    # Key Numbers
    # ---------------------------
    doc.add_heading("Key Numbers (what the engine is looking at)", level=1)
    if key_nums:
        doc.add_paragraph(f"Revenue (TTM): {fmt_money_bn(key_nums.get('Revenue (TTM)'))}", style="List Bullet")
        doc.add_paragraph(f"Free Cash Flow (TTM): {fmt_money_bn(key_nums.get('Free Cash Flow (TTM)'))}", style="List Bullet")
        doc.add_paragraph(f"FCF Margin (TTM): {fmt_pct(key_nums.get('FCF Margin (TTM %)'))}", style="List Bullet")
        doc.add_paragraph(f"Revenue Growth (YoY): {fmt_pct(key_nums.get('Revenue Growth (TTM YoY %)'))}", style="List Bullet")
        doc.add_paragraph(f"FCF Growth (YoY): {fmt_pct(key_nums.get('FCF Growth (TTM YoY %)'))}", style="List Bullet")
        doc.add_paragraph(f"Market Cap: {fmt_money_bn(key_nums.get('Market Cap'))}", style="List Bullet")
        doc.add_paragraph(f"FCF Yield: {fmt_ratio(key_nums.get('FCF Yield'))} (higher usually = cheaper vs cash)", style="List Bullet")
        doc.add_paragraph(f"Net Debt: {fmt_money_bn(key_nums.get('Net Debt'))}", style="List Bullet")
        doc.add_paragraph(f"Net Debt / FCF: {fmt_ratio(key_nums.get('Net Debt / FCF'))} (lower usually = safer)", style="List Bullet")
    else:
        doc.add_paragraph("Key numbers not available in this run.")

    # ---------------------------
    # How score works
    # ---------------------------
    doc.add_heading("How the Score Works (0–100)", level=1)
    doc.add_paragraph(
        "Think of the score like a report card. We split the decision into five buckets and add them up."
    )
    doc.add_paragraph("Buckets:", style="List Bullet")
    doc.add_paragraph("Cash Level: how much real cash the business produces.", style="List Bullet 2")
    doc.add_paragraph("Valuation: whether the stock looks expensive vs the cash it produces.", style="List Bullet 2")
    doc.add_paragraph("Growth: whether the business is accelerating or slowing.", style="List Bullet 2")
    doc.add_paragraph("Quality: how efficiently revenue becomes free cash flow (margin).", style="List Bullet 2")
    doc.add_paragraph("Balance/Risk: debt + recent negative news risk.", style="List Bullet 2")
    doc.add_paragraph(f"Bucket scores today: {buckets}")

    # ---------------------------
    # Trace (how we arrive)
    # ---------------------------
    doc.add_heading("How we arrived at the score (step-by-step trace)", level=1)
    doc.add_paragraph("The engine follows the same sequence every run:")

    doc.add_paragraph("Step 1 — Collect raw inputs", style="List Bullet")
    doc.add_paragraph("Fundamentals: revenue, operating cash flow, capex → to compute free cash flow.", style="List Bullet 2")
    doc.add_paragraph("Market: price + market cap → to convert cash into valuation.", style="List Bullet 2")
    doc.add_paragraph("Peers: LYFT, DASH → sanity-check relative strength.", style="List Bullet 2")
    doc.add_paragraph("News: SEC + Finnhub → detect repeating risk themes and severe negatives.", style="List Bullet 2")

    doc.add_paragraph("Step 2 — Convert inputs into comparable metrics", style="List Bullet")
    doc.add_paragraph("TTM Revenue: sales in last 12 months.", style="List Bullet 2")
    doc.add_paragraph("TTM Free Cash Flow: cash left after operations + capex in last 12 months.", style="List Bullet 2")
    doc.add_paragraph("FCF Margin: % of revenue that becomes free cash.", style="List Bullet 2")
    doc.add_paragraph("FCF Yield: FCF divided by market cap (cash return vs price).", style="List Bullet 2")
    doc.add_paragraph("Net Debt / FCF: leverage burden relative to cash generation.", style="List Bullet 2")

    doc.add_paragraph("Step 3 — Score each bucket", style="List Bullet")
    doc.add_paragraph("Buckets reward strong cash, reasonable price, healthy growth, and good margins.", style="List Bullet 2")
    doc.add_paragraph("Risk bucket deducts points when leverage is heavy or negative news repeats.", style="List Bullet 2")

    doc.add_paragraph("Step 4 — Add buckets → Final score + rating", style="List Bullet")
    doc.add_paragraph("BUY = strong overall signal; HOLD = mixed; AVOID = weak or risky.", style="List Bullet 2")

    # ---------------------------
    # Bucket explanations
    # ---------------------------
    doc.add_heading("Bucket Explanations (why it scored this way)", level=1)

    explanations = [
        ("Cash Level", f"{buckets.get('cash_level')} points",
         "This checks whether the company produces real cash (free cash flow). Cash matters because it funds growth without borrowing.",
         "Good if: FCF is positive and rising. Bad if: FCF is negative or unstable."),
        ("Valuation", f"{buckets.get('valuation')} points",
         "This checks whether you’re paying a reasonable price for that cash. If price rises faster than cash, valuation becomes stretched.",
         "Good if: FCF Yield is healthy. Bad if: yield is tiny (meaning expensive vs cash)."),
        ("Growth", f"{buckets.get('growth')} points",
         "This checks whether the business is expanding. Growth supports higher valuations—slowing growth can hurt even profitable companies.",
         "Good if: Revenue and FCF growth are strong. Bad if: growth slows sharply or turns negative."),
        ("Quality", f"{buckets.get('quality')} points",
         "This checks how efficiently revenue turns into free cash flow (FCF margin). Higher margin = better economics.",
         "Good if: FCF margin expands. Bad if: margins compress."),
        ("Balance/Risk", f"{buckets.get('balance_risk')} points",
         "This is the safety bucket: leverage (net debt) plus repeating negative news themes (regulatory / insurance / labor / safety).",
         "Good if: debt is manageable and negatives are rare. Bad if: risk tags repeat and shocks worsen."),
    ]

    for title, pts, meaning, goodbad in explanations:
        doc.add_paragraph(f"{title} — {pts}", style="Heading 2")
        doc.add_paragraph(meaning)
        doc.add_paragraph("How to interpret:", style="List Bullet")
        doc.add_paragraph(goodbad, style="List Bullet 2")

        if title == "Balance/Risk":
            neg_7d = int(news_summary.get("neg_7d", 0) or 0)
            shock_7d = int(news_summary.get("shock_7d", 0) or 0)
            doc.add_paragraph("Recent risk snapshot:", style="List Bullet")
            doc.add_paragraph(f"Negative headlines (7d): {neg_7d}", style="List Bullet 2")
            doc.add_paragraph(f"Shock score (7d): {shock_7d} (more negative = worse)", style="List Bullet 2")

            top_neg = bucket_evidence.get("top_negative_headlines_7d", []) or []
            if top_neg:
                doc.add_paragraph("Example negative headlines (7d):", style="List Bullet")
                for h in top_neg[:5]:
                    doc.add_paragraph(str(h), style="List Bullet 2")

    # ---------------------------
    # Metric cheat sheet (good/ok/bad)
    # ---------------------------
    doc.add_heading("Metric Cheat Sheet (what’s good vs bad)", level=1)
    doc.add_paragraph(
        "Use this as a translation key. These are rules of thumb to help non-finance readers interpret the numbers."
    )

    mcs = metric_cheat_sheet()
    t = doc.add_table(rows=1, cols=len(mcs.columns))
    for i, c in enumerate(mcs.columns):
        t.rows[0].cells[i].text = str(c)
    for _, row in mcs.iterrows():
        cells = t.add_row().cells
        for i, c in enumerate(mcs.columns):
            cells[i].text = str(row.get(c, ""))
    _format_table(t, header_fill="E8EEF7", header_font_size=10, body_font_size=9)

    # ---------------------------
    # Veracity
    # ---------------------------
    doc.add_heading("Veracity (can you verify the sources?)", level=1)
    doc.add_paragraph(
        "This measures whether the evidence is easy to verify (source quality + URLs). "
        "It does NOT predict the stock price."
    )
    doc.add_paragraph(f"Confidence score: {confidence_score}/100", style="List Bullet")
    for r in confidence_reasons:
        doc.add_paragraph(str(r), style="List Bullet 2")

    doc.add_paragraph("Key veracity stats:", style="List Bullet")
    doc.add_paragraph(f"Total news items analyzed: {veracity.get('total_rows')}", style="List Bullet 2")
    doc.add_paragraph(f"Top-tier share: {veracity.get('top_tier_pct')}%", style="List Bullet 2")
    doc.add_paragraph(f"Redirect-like URLs: {veracity.get('redirect_like_rows')}", style="List Bullet 2")
    doc.add_paragraph(f"Duplicate headlines (estimate): {veracity.get('duplicate_titles_est')}", style="List Bullet 2")
    doc.add_paragraph("You can edit what counts as TOP tier in export/source_whitelist.csv.")

    # ---------------------------
    # NEXT STEPS (painless, recipe-style)
    # ---------------------------
    doc.add_heading("Next Steps (exactly what to do)", level=1)
    doc.add_paragraph("This is the least painful way to validate the decision.")

    doc.add_paragraph("✅ 3-minute validation (minimum recommended):", style="List Bullet")
    doc.add_paragraph("1) Read Executive Summary + Red Flags.", style="List Bullet 2")
    doc.add_paragraph("2) Read the Curated Evidence Pack (10–12 items) below.", style="List Bullet 2")
    doc.add_paragraph("3) Decide: are the negatives repeating (pattern) or isolated (noise)?", style="List Bullet 2")

    doc.add_paragraph("✅ 10-minute validation (better confidence):", style="List Bullet")
    doc.add_paragraph("4) Read Bucket Explanations — each explains what drives points.", style="List Bullet 2")
    doc.add_paragraph("5) Use the Metric Cheat Sheet to label numbers as good/ok/bad.", style="List Bullet 2")
    doc.add_paragraph("6) Compare UBER vs peers in the Excel ‘Peers’ tab.", style="List Bullet 2")

    doc.add_paragraph("✅ 30-minute diligence (deep dive):", style="List Bullet")
    doc.add_paragraph("7) Review the Risk Dashboard (Appendix): focus on repeating tags.", style="List Bullet 2")
    doc.add_paragraph("8) Open the HTML report and click the most severe negatives to verify details.", style="List Bullet 2")
    doc.add_paragraph("9) If risks are repeating AND severe, manually downgrade your confidence.", style="List Bullet 2")

    doc.add_paragraph(
        "If you do only ONE thing: verify the worst negative headline(s) and see if they repeat. "
        "That’s usually where surprises live."
    )

    # ---------------------------
    # Curated evidence pack (Word-friendly)
    # ---------------------------
    doc.add_heading("Curated Evidence Pack (what to verify first)", level=1)
    doc.add_paragraph(
        "These are the highest-priority items selected by severity + risk tag + recency. "
        "The HTML report is best for clicking; below we include titles + URLs so you can verify."
    )

    if curated_pack is None or curated_pack.empty:
        doc.add_paragraph("No curated evidence available.")
    else:
        for _, row in curated_pack.iterrows():
            line = f"{row.get('published_at')} [{row.get('risk_tag')}] ({row.get('impact_score')}): {row.get('title')}"
            doc.add_paragraph(line, style="List Bullet")
            doc.add_paragraph(str(row.get("url", "")), style="List Bullet 2")

    # ---------------------------
    # Appendix (tables go here)
    # ---------------------------
    doc.add_page_break()
    doc.add_heading("Appendix (details)", level=1)

    doc.add_heading("A1) Peer Comps (selected columns)", level=2)
    if comps is not None and not comps.empty:
        view = comps.copy()
        keep = [c for c in [
            "ticker", "price", "market_cap_usd_bn",
            "revenue_ttm_usd_bn", "revenue_ttm_yoy_pct",
            "fcf_ttm_usd_bn", "fcf_ttm_yoy_pct",
            "fcf_margin_ttm_pct", "fcf_yield",
            "net_debt_usd_bn", "net_debt_to_fcf_ttm",
        ] if c in view.columns]
        if keep:
            view = view[keep]

        table = doc.add_table(rows=1, cols=len(view.columns))
        hdr = table.rows[0].cells
        for i, c in enumerate(view.columns):
            hdr[i].text = str(c)
        for _, row in view.iterrows():
            cells = table.add_row().cells
            for i, c in enumerate(view.columns):
                cells[i].text = str(row.get(c, ""))
        _format_table(table, header_fill="EDEDED", header_font_size=10, body_font_size=9)
    else:
        doc.add_paragraph("No comps data available.")

    doc.add_heading("A2) Risk Dashboard (first 40 rows)", level=2)
    if risk_dash is not None and not risk_dash.empty:
        rd = risk_dash.head(40).copy()
        table = doc.add_table(rows=1, cols=len(rd.columns))
        hdr = table.rows[0].cells
        for i, c in enumerate(rd.columns):
            hdr[i].text = str(c)
        for _, row in rd.iterrows():
            cells = table.add_row().cells
            for i, c in enumerate(rd.columns):
                cells[i].text = str(row.get(c, ""))
        _format_table(table, header_fill="EDEDED", header_font_size=10, body_font_size=9)
    else:
        doc.add_paragraph("No risk dashboard available.")

    doc.add_heading("A3) Evidence links (Top-tier, first 40)", level=2)
    doc.add_paragraph("Tip: HTML report is easiest to click. This appendix preserves the links in text form.")
    if evidence_top is not None and not evidence_top.empty:
        for _, row in evidence_top.head(40).iterrows():
            doc.add_paragraph(
                f"{row.get('published_at')} [{row.get('risk_tag')}] ({row.get('impact_score')}): {row.get('title')}",
                style="List Bullet",
            )
            doc.add_paragraph(str(row.get("url", "")), style="List Bullet 2")
    else:
        doc.add_paragraph("No top-tier evidence available.")

    doc.save(out_path)


# ---------------------------
# Excel formatting helpers
# ---------------------------
def _autosize_columns(ws):
    for col in ws.columns:
        max_len = 10
        col_letter = col[0].column_letter
        for cell in col:
            try:
                v = "" if cell.value is None else str(cell.value)
                max_len = max(max_len, len(v))
            except Exception:
                pass
        ws.column_dimensions[col_letter].width = min(60, max_len + 2)


def _style_table(ws, name: str):
    try:
        tab = Table(displayName=re.sub(r"[^A-Za-z0-9]", "", name)[:25] or "Table1", ref=ws.dimensions)
        style = TableStyleInfo(name="TableStyleMedium9", showRowStripes=True, showColumnStripes=False)
        tab.tableStyleInfo = style
        ws.add_table(tab)
    except Exception:
        pass


def _apply_number_formats(ws, header_row: int = 1):
    headers = [c.value for c in ws[header_row]]
    col_map = {str(h): i + 1 for i, h in enumerate(headers) if h is not None}

    def set_fmt(col_name: str, fmt: str):
        if col_name not in col_map:
            return
        j = col_map[col_name]
        for r in range(header_row + 1, ws.max_row + 1):
            ws.cell(row=r, column=j).number_format = fmt

    # Raw dollars (big integers)
    for c in ["market_cap", "revenue_ttm", "fcf_ttm", "cash", "debt", "net_debt"]:
        set_fmt(c, "#,##0")

    # $ billions columns
    for c in ["market_cap_usd_bn", "revenue_ttm_usd_bn", "fcf_ttm_usd_bn", "cash_usd_bn", "debt_usd_bn", "net_debt_usd_bn"]:
        set_fmt(c, "0.00")

    # Percent-like columns (these are already percent numbers, not 0-1 fractions)
    for c in ["revenue_ttm_yoy_pct", "fcf_ttm_yoy_pct", "fcf_margin_ttm_pct"]:
        set_fmt(c, "0.00")

    set_fmt("fcf_yield", "0.0000")
    set_fmt("net_debt_to_fcf_ttm", "0.00")
    set_fmt("price", "0.00")

    # counts / impact
    for c in ["impact_score", "neg_count_7d", "neg_count_30d", "shock_7d", "shock_30d", "proxy_score_7d", "proxy_score_30d"]:
        set_fmt(c, "0")


def _add_df_sheet(wb: Workbook, name: str, df: pd.DataFrame, apply_formats: bool = True):
    ws = wb.create_sheet(title=name)
    if df is None or df.empty:
        ws["A1"] = "No data"
        return ws

    for r in dataframe_to_rows(df, index=False, header=True):
        ws.append(r)

    header_font = Font(bold=True)
    for cell in ws[1]:
        cell.font = header_font
        cell.alignment = Alignment(wrap_text=True, vertical="top")

    ws.freeze_panes = "A2"
    _style_table(ws, name)
    if apply_formats:
        _apply_number_formats(ws)
    _autosize_columns(ws)
    return ws


def write_excel_report(
    out_path: Path,
    decision_summary: dict,
    comps: pd.DataFrame,
    proxy: pd.DataFrame,
    risk_dash: pd.DataFrame,
    whitelist: pd.DataFrame,
    veracity: dict,
    confidence_score: int,
    confidence_reasons: List[str],
    bucket_evidence: dict,
    curated_pack: pd.DataFrame,
    evidence_top: pd.DataFrame,
    evidence_all: pd.DataFrame,
):
    out_path.parent.mkdir(parents=True, exist_ok=True)
    wb = Workbook()
    wb.remove(wb.active)

    ws_sum = wb.create_sheet("Summary")
    ws_sum.append(["key", "value"])
    rows = [
        ("ticker", decision_summary.get("ticker", TICKER)),
        ("as_of", decision_summary.get("as_of", "")),
        ("score", decision_summary.get("score", "")),
        ("rating", decision_summary.get("rating", "")),
        ("confidence_score", confidence_score),
        ("confidence_reasons", json.dumps(confidence_reasons)),
        ("bucket_scores", json.dumps(decision_summary.get("bucket_scores", {}))),
        ("red_flags", json.dumps(decision_summary.get("red_flags", []))),
        ("news_summary", json.dumps(decision_summary.get("news_summary", {}))),
        ("news_sentiment_proxy", json.dumps(decision_summary.get("news_sentiment_proxy", {}))),
    ]
    for k, v in rows:
        ws_sum.append([k, v])

    for cell in ws_sum[1]:
        cell.font = Font(bold=True)
    ws_sum.freeze_panes = "A2"
    _autosize_columns(ws_sum)

    _add_df_sheet(wb, "SourceWhitelist", whitelist, apply_formats=False)
    _add_df_sheet(wb, "Peers", comps, apply_formats=True)
    _add_df_sheet(wb, "News_Proxy", proxy, apply_formats=True)
    _add_df_sheet(wb, "Risk_Dashboard", risk_dash, apply_formats=True)
    _add_df_sheet(wb, "Curated_Evidence", curated_pack, apply_formats=False)
    _add_df_sheet(wb, "Evidence_TopTier", evidence_top, apply_formats=False)
    _add_df_sheet(wb, "Evidence_All", evidence_all, apply_formats=False)

    ver_df = pd.DataFrame(
        [{"metric": k, "value": json.dumps(v) if isinstance(v, (dict, list)) else v} for k, v in veracity.items()]
    )
    _add_df_sheet(wb, "Veracity", ver_df, apply_formats=False)

    bucket_df = pd.DataFrame(
        [{"section": k, "details": json.dumps(v)} for k, v in bucket_evidence.items()]
    )
    _add_df_sheet(wb, "BucketEvidence", bucket_df, apply_formats=False)

    wb.save(out_path)


def main():
    ensure_dirs()

    decision_summary = load_json(OUTPUTS / "decision_summary.json")
    news_df = safe_read_csv(DATA_PROCESSED / "news_unified.csv")
    comps = safe_read_csv(DATA_PROCESSED / "comps_snapshot.csv")
    proxy = safe_read_csv(DATA_PROCESSED / "news_sentiment_proxy.csv")
    risk_dash = safe_read_csv(DATA_PROCESSED / "news_risk_dashboard.csv")

    whitelist = load_source_whitelist()
    veracity = build_veracity_summary(news_df, whitelist)
    confidence_score, confidence_reasons = compute_confidence_score(veracity)

    evidence_top, evidence_all = build_evidence_tables(news_df, ticker=TICKER, whitelist=whitelist, days=30, max_rows=350)
    bucket_evidence = build_bucket_evidence(decision_summary)

    comps_h = add_human_units_comps(comps)

    # Curated pack: prefer TOP tier evidence if we have enough; else fall back to all
    base_for_curate = evidence_top if evidence_top is not None and len(evidence_top) >= 8 else evidence_all
    curated_pack = curate_evidence(base_for_curate, top_n=12)

    (OUTPUTS / f"veracity_{TICKER}.json").write_text(json.dumps(veracity, indent=2), encoding="utf-8")

    html_path = OUTPUTS / f"decision_report_{TICKER}.html"
    write_html_report(
        html_path,
        title=f"Decision Report — {TICKER}",
        decision_summary=decision_summary,
        comps=comps_h,
        risk_dash=risk_dash,
        veracity=veracity,
        confidence_score=confidence_score,
        confidence_reasons=confidence_reasons,
        curated_pack=curated_pack,
        evidence_top=evidence_top,
        evidence_all=evidence_all,
    )

    docx_path = EXPORT / f"{TICKER}_Investment_Report.docx"
    write_word_report(
        docx_path,
        decision_summary=decision_summary,
        comps=comps_h,
        proxy=proxy,
        risk_dash=risk_dash,
        veracity=veracity,
        confidence_score=confidence_score,
        confidence_reasons=confidence_reasons,
        bucket_evidence=bucket_evidence,
        curated_pack=curated_pack,
        evidence_top=evidence_top,
    )

    xlsx_path = EXPORT / f"{TICKER}_Investment_Report.xlsx"
    write_excel_report(
        xlsx_path,
        decision_summary=decision_summary,
        comps=comps_h,
        proxy=proxy,
        risk_dash=risk_dash,
        whitelist=whitelist,
        veracity=veracity,
        confidence_score=confidence_score,
        confidence_reasons=confidence_reasons,
        bucket_evidence=bucket_evidence,
        curated_pack=curated_pack,
        evidence_top=evidence_top,
        evidence_all=evidence_all,
    )

    print("DONE ✅ Reports created:")
    print(f"- {html_path}")
    print(f"- {OUTPUTS / f'veracity_{TICKER}.json'}")
    print(f"- {docx_path}")
    print(f"- {xlsx_path}")
    print(f"- Whitelist you can edit: {WHITELIST_PATH}")


if __name__ == "__main__":
    main()
